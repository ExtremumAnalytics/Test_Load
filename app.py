import os
import time
import base64
import json
from io import BytesIO
import re

import docx
import asyncio
import pytz
from werkzeug.utils import secure_filename

from azure.search.documents._generated.models import SearchMode, VectorizedQuery
from azure.search.documents.indexes._generated.models import SemanticConfiguration, SemanticPrioritizedFields, \
    SemanticField, SemanticSearch, ScoringProfile, TextWeights
from flask import Flask, jsonify, url_for, flash
from flask import render_template, request, g, redirect, session
import tempfile

# For Virtual Analyst DB integration
from sqlalchemy import create_engine, inspect

# sentiment and word cloud use only
import nltk
from nltk.tokenize import word_tokenize
from wordcloud import WordCloud
from nltk.sentiment.vader import SentimentIntensityAnalyzer
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation
from nltk.corpus import stopwords

# for azure credentials use only
from azure.storage.blob import BlobServiceClient, ContentSettings
from azure.cognitiveservices.vision.computervision.models import OperationStatusCodes

# langchain module
from langchain.prompts import PromptTemplate
from langchain.chains.summarize import load_summarize_chain
from langchain_community.document_loaders import Docx2txtLoader, \
    UnstructuredExcelLoader, \
    AssemblyAIAudioTranscriptLoader
from langchain_community.document_loaders import WebBaseLoader
from langchain_community.chat_models import AzureChatOpenAI
from langchain_community.vectorstores import AzureSearch
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.chains.llm import LLMChain
from langchain.schema import Document
from langchain_community.utilities import DuckDuckGoSearchAPIWrapper
from langchain_community.tools import DuckDuckGoSearchRun
from langchain.schema import AIMessage
from langchain_core.runnables import RunnablePassthrough
# for azure vector db index creation
from azure.search.documents import SearchClient
from langchain_community.document_loaders import PyPDFLoader
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    SearchFieldDataType,
    SimpleField,
    SearchableField,
    SearchIndex,
    SearchField,
    VectorSearch,
    VectorSearchProfile,
    HnswAlgorithmConfiguration
)

# for mp3 to pdf
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# for webscraping
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse

# for EDA
import pandas as pd
import matplotlib
from pandasai import Agent
from pandasai.llm import AzureOpenAI

from datetime import datetime

# Socket IO
from flask_socketio import SocketIO, emit

# For database connection
import datetime
import mysql.connector
# from pymongo import MongoClient
import pymongo
import io

# Import folder file
from EA.Database.database import db, create_all_tables
from EA.Database.database import UserRole, UserDetails, ChatHistory, DatabaseDetailsSave, SummaryHistory
from EA.Config.configuration import (blob_service_client, container_client, computervision_client,
                                     embeddings, vector_store_address, vector_store_password, container_name,
                                     main_key, DB_USERNAME, DB_PASSWORD, DB_HOST, DB_NAME)
from EA.Logger.logging import setup_database_logger, CustomLoggerAdapter
from EA.Interrupt.flag import check_stop_flag, write_stop_flag_to_csv, read_stop_flag_from_csv
from EA.Error_Handling.checkError import check_file, check_error

chunk_size = 8000
chunk_overlap = 400
custom_prompt = ''
chain_type = 'map_reduce'
num_summaries = 1

matplotlib.use('Agg')

blob_list_length = 0
tot_file = 0
tot_suc = 0
tot_fail = 0
# word cloud
text_word_cloud = ''

# nltk library download
nltk.download('vader_lexicon')
nltk.download('stopwords')
nltk.download('punkt')

global blob_client, logger, chat_history_list, Limit_By_Size, mb_pop, file_size_bytes, df, png_file, loader

app = Flask(__name__, template_folder="Templates")
app.debug = True
app.static_folder = "static"
app.config["SQLALCHEMY_DATABASE_URI"] = f"mysql+pymysql://{DB_USERNAME}:{DB_PASSWORD}@{DB_HOST}/{DB_NAME}"
db.init_app(app)

# Create tables
create_all_tables(app)
app.config['DEBUG'] = True
app.secret_key = os.urandom(24)
socketio = SocketIO(app, manage_session=True)

env_mapping_dict = {"http://cognilink-prod.azurewebsites.net/": "prod",
                    "http://cognilink-dev.azurewebsites.net/": "dev", "http://127.0.0.1:5000/": "dev"}

# Define timezones
utc_zone = pytz.utc
ist_zone = pytz.timezone('Asia/Kolkata')


def convert_to_ist(date):
    # Ensure the date is timezone aware, if it is naive, set it to UTC first
    if date.tzinfo is None:
        date = utc_zone.localize(date)
    # Convert to IST
    date_ist = date.astimezone(ist_zone)
    # Convert the date object to an ISO 8601 string
    return date_ist.isoformat()


def set_model():
    model = session.get('engine', 'gpt-4-0125-preview')  # Default to 'gpt-4-0125-preview'
    if model == "GPT-3.5 turbo":
        deployment_name = "gpt-35-turbo"
    elif model == "GPT-4":
        deployment_name = "gpt-4-0125-preview"
    elif model == "GPT-4o":
        deployment_name = ""  # upcoming
    else:
        deployment_name = "gpt-4-0125-preview"
    return deployment_name


def create_or_pass_folder(container_client_, session):
    """
     Creates a folder in the specified Azure Blob Storage container if it does not already exist.

     Args:
         container_client_ (BlobServiceClient): Azure Blob Service client instance.
         session (dict): User session dictionary containing the 'login_pin'.

     Returns:
         str: Message indicating the result of the folder creation process.

     Side Effects:
         Sets the global variable `g.flag` to 1 on success and 0 on failure.
         Logs the success or failure of the folder creation process.
     """

    if 'login_pin' in session:
        user_folder = str(session['login_pin'])
        try:
            container_client_.get_blob_client(container_name, user_folder).upload_blob("")
            g.flag = 1  # Set flag to 1 on success1
            logger.info(f"Function create_or_pass_folder successfully created folder")
            return f"Folder '{user_folder}' successfully created."
        except Exception as e:
            g.flag = 0  # Set flag to 1 on success1
            logger.error(f"Function create_or_pass_folder error", exc_info=True)
            if "BlobNotFound" in str(e):
                try:
                    container_client_.get_blob_client(container_name, user_folder).create_container()
                    return f"Folder '{user_folder}' successfully created."
                except Exception as e:
                    return f"Error creating folder '{user_folder}': {str(e)}"
            else:
                g.flag = 0
                logger.error(f"Function create_or_pass_folder Error creating folder'{user_folder}': {str(e)} ",
                             exc_info=True)
                return f"Error creating folder '{user_folder}': {str(e)}"
    else:
        g.flag = 0
        logger.error(f"Function create_or_pass_folder login_pin' not found in session. ", exc_info=True)
        return "Error: 'login_pin' not found in session."


def clean_filename(filename):
    # Remove special characters and replace spaces with underscores
    cleaned_filename = re.sub(r'[^\w\s.-]', '', filename)
    cleaned_filename = re.sub(r'\s+', '_', cleaned_filename)
    return cleaned_filename


def upload_to_blob(file_content, session, blob_service_client_, container_name_):
    """Uploads a file to Azure Blob Storage with enhanced security and error handling.

    Args:
        file_content (http.MultipartFile): The file object to upload.
        session (dict): User session dictionary.
        blob_service_client_ (BlobServiceClient): Azure Blob Service client instance.
        container_name_ (str): Name of the Azure Blob Storage container.

    Returns:
        str: URL of the uploaded blob or an error message.
    """
    global blob_client

    if 'login_pin' not in session:
        return "Error: 'login_pin' not found in session."

    try:
        folder_name = f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}"

        # Clean the file name
        cleaned_filename = clean_filename(file_content.filename)

        blob_name = f"{folder_name}/{cleaned_filename}"
        blob_client = blob_service_client_.get_blob_client(container=container_name_, blob=blob_name)

        # Read the content of file_content
        content = file_content.read()

        # Determine if the file is an image based on its extension
        file_extension = cleaned_filename.split('.')[-1].lower()
        if file_extension in ['jpg', 'jpeg', 'png']:
            # Reset the file pointer for the image
            file_content.seek(0)
            content_type = file_content.mimetype or 'application/octet-stream'
            blob_client.upload_blob(file_content, blob_type="BlockBlob", overwrite=True,
                                    content_settings=ContentSettings(content_type=content_type))
        else:
            # Upload the non-image content to Azure Blob Storage
            blob_client.upload_blob(content, blob_type="BlockBlob",
                                    content_settings=ContentSettings(content_type=file_content.content_type),
                                    overwrite=True)

        g.flag = 1  # Set flag to 1 on success
        logger.info("Function upload_to_blob successfully uploaded the blob")

        g.flag = 1  # Set flag to 1 on success
        logger.info("Function upload_to_blob successfully uploaded the blob")

        return blob_client.url

    except Exception as e:
        g.flag = 0  # Set flag to 0 on error
        logger.error("Function upload_to_blob error", exc_info=True)
        return f"Error: {str(e)}"


def update_bar_chart_from_blob(session, blob_service_client_, container_name_):
    """
    Updates a bar chart in the user session based on the types of files in a specified Azure Blob Storage folder.

    Args:
        session (dict): User session dictionary containing the 'login_pin'.
        blob_service_client_ (BlobServiceClient): Azure Blob Service client instance.
        container_name_ (str): Name of the Azure Blob Storage container.

    Returns:
        list: List of blobs in the specified folder.

    Side Effects:
        Updates the 'bar_chart_ss' in the session with the count of different file types.
        Emits a 'update_bar_chart' event via socketio to notify clients about the updated bar chart.
        Sets the global variable `g.flag` to 1 on success and 0 on failure.
        Logs the success or failure of the bar chart update process.
    """
    bar_chart = {}
    blob_list = []
    try:
        # Get a list of blobs in the specified folder
        blob_list = blob_service_client_.get_container_client(container_name_).list_blobs(
            name_starts_with=f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}")

        # Iterate through each blob in the folder
        for blob in blob_list:
            if blob.name.split('/')[2] != 'draft':
                file_name = blob.name.split('/')[-1]  # Extract file name from blob path
                file_name = file_name.lower()
                # Update the bar_chart dictionary based on file type
                if file_name.endswith('.pdf'):
                    file_type = 'PDF'
                elif file_name.endswith('.docx') or file_name.endswith('.doc'):
                    file_type = 'DOCX'
                elif file_name.endswith('.csv'):
                    file_type = 'CSV'
                elif file_name.endswith('.mp3'):
                    file_type = 'MP3'
                elif file_name.endswith('.xlsx') or file_name.endswith('.xlscd') or file_name.endswith('.xls'):
                    file_type = 'XLSX'
                elif file_name.endswith('.jpg') or file_name.endswith('.png') or file_name.endswith('.jpeg'):
                    file_type = 'Image'
                elif 'https://' or 'http://' in file_name:
                    file_type = 'Website'
                else:
                    file_type = 'Other'

                # Update bar_chart dictionary
                if file_type in bar_chart:
                    bar_chart[file_type] += 1
                else:
                    bar_chart[file_type] = 1
        # Emit an event to notify clients about the updated bar chart
        socketio.emit('update_bar_chart', {
            'labels': list(bar_chart.keys()),
            'values': list(bar_chart.values()),
            'pin': session['login_pin']
        })
        g.flag = 1  # Set flag to 1 on success1
        logger.info(f"Function update_bar_chart_from_blob successfully Return blob_list")
    except Exception as e:
        g.flag = 0  # Set flag to 1 on success1
        logger.error(f"Function update_bar_chart_from_blob error", exc_info=True)
        print('blob_list error------>', str(e))

    # Return the updated bar_chart dictionary
    return blob_list


def update_when_file_delete():
    """
    Updates session and charts when a file is deleted from Azure Blob Storage.

    Side Effects:
        Sets the global variables related to blob list, source URL, loader, total files, bar chart URL, and txt to PDF.
        Updates session counts for various metrics.
        Logs the success or failure of the file update process.
        Emits events to update charts on the client side.

    Returns:
        Response: JSON response indicating the result of the operation.
    """
    global blob_list_length, loader, tot_file
    blob_list_length = 0
    tot_suc = 0
    tot_fail = 0

    start_time = time.time()
    # Lists to store progress of files loading
    session['embedding_not_created'] = []
    session['failed_files'] = []
    session['progress_files'] = []
    all_blobs = blob_service_client.get_container_client(container_name).list_blobs(
        name_starts_with=f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}")
    all_blobs_list = list(all_blobs)  # Convert to list to enable filtering

    blob_list = [blob for blob in all_blobs_list if not (
            blob.name.endswith('.csv') or
            blob.name.endswith('.CSV') or
            blob.name.endswith('.jpg') or
            blob.name.endswith('.png') or
            blob.name.endswith('.text') or
            blob.name.endswith('.jpeg') or
            blob.name.endswith('.mp3') or
            blob.name.endswith('.json') or
            blob.name.endswith('_schema.xls') or
            blob.name.endswith('_schema.xlsx')
    )]

    # Update session counts for CSV files directly
    csv_files_count = len(all_blobs_list) - len(blob_list)
    tot_suc += csv_files_count

    blob_list_jpg = [blob for blob in blob_list if not (
            blob.name.endswith('.jpg') or blob.name.endswith('.JPG') or blob.name.endswith(
        '.PNG') or blob.name.endswith('.png') or blob.name.endswith(
        '.jpeg') or blob.name.endswith('.JPEG'))]

    # Initialize SearchClient
    search_client = SearchClient(
        endpoint=vector_store_address,
        index_name=f"cognilink-{session['env_map']}-{session['login_pin']}",
        credential=AzureKeyCredential(vector_store_password)
    )
    results = search_client.search(search_text="*", select="*", include_total_count=True)

    VecTor_liSt = []

    unique_documents = set()

    for result in results:
        embeddings_dict = json.loads(result['metadata'])
        document = embeddings_dict.get('documents')
        if document and document not in unique_documents:
            VecTor_liSt.append(document)
            unique_documents.add(document)

    # Filtering out blobs whose names are in VecTor_liSt
    result_loop = [blob for blob in blob_list_jpg if
                   not any(blob['name'].endswith(item) for item in VecTor_liSt)]

    # Further filtering based on VecTor_liSt
    result_loop1 = []
    for blob in result_loop:
        # Extract the URL part for comparison
        if 'https://' in blob['name'] or 'http://' in blob['name']:
            url_part = blob['name'].split(f"{session['login_pin']}/")[1]
        else:
            url_part = blob['name']

        # Append to result_loop if the URL part is not in VecTor_liSt
        if url_part not in VecTor_liSt:
            result_loop1.append(blob)

    blob_list_length = len(result_loop1)
    if blob_list_length == 0:
        session['bar_chart_ss'] = {}
        session['over_all_readiness'] = 0
        session['total_success_rate'] = 0
        session['total_files_list'] = 0
        session['successful_list'] = 0
        session['failed_list'] = 0
        session['progress_list'] = 0
        print("No data Load in storage")
        return jsonify({'message': 'No new file uploaded, old one Loaded or upload new file.'})
    session['total_files_list'] = blob_list_length + csv_files_count
    socketio.emit('progress', {'percentage': 20, 'pin': session['login_pin']})
    try:
        if blob_list_length != 0:
            for blob in result_loop1:
                final_chunks = []

                if check_stop_flag():
                    write_stop_flag_to_csv(session['login_pin'], 'False')
                    print("Data Load Cancelled")
                    break
                if 'https://' or 'http://' in blob.name:
                    # print("name------>URL", blob.name)
                    url_part = blob.name.split(str(session['login_pin']) + '/')[1]
                    # print("name------>URL", url_part)
                    loader = WebBaseLoader(url_part)
                    session['embedding_not_created'].append(url_part)
                    socketio.emit('pending', session['embedding_not_created'])

                blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob.name)
                file_name = blob.name.split('/')[-1]  # Extract file name from blob path
                if file_name.endswith('.pdf') or file_name.endswith('.PDF') or file_name.endswith('.mp3'):
                    temp_path = blob_client.url
                elif file_name.endswith('.xls') or file_name.endswith('.xlsx') or file_name.endswith(
                        '.docx') or file_name.endswith('.doc'):
                    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                        temp_path = temp_file.name
                        blob_data = blob_client.download_blob()
                        blob_data.readinto(temp_file)

                if '.pdf' in file_name or '.PDF' in file_name:
                    loader = PyPDFLoader(temp_path)
                    session['embedding_not_created'].append(file_name)
                    socketio.emit('pending', session['embedding_not_created'])

                # elif '.mp3' in file_name:
                #     loader = AssemblyAIAudioTranscriptLoader(temp_path, api_key="5bbe5761b36b4ff885dbd18836c3c723")
                #     session['embedding_not_created'].append(file_name)
                #     socketio.emit('pending', session['embedding_not_created'])

                elif '.docx' in file_name or '.doc' in file_name:
                    loader = Docx2txtLoader(temp_path)
                    session['embedding_not_created'].append(file_name)
                    socketio.emit('pending', session['embedding_not_created'])

                elif '.xlsx' in file_name or '.xls' in file_name:
                    loader = UnstructuredExcelLoader(temp_path)
                    session['embedding_not_created'].append(file_name)
                    socketio.emit('pending', session['embedding_not_created'])

                if check_stop_flag():
                    write_stop_flag_to_csv(session['login_pin'], 'False')
                    print("Data Load Cancelled")
                    break

                s_url = 'https://testcongnilink.blob.core.windows.net/congnilink-container/' + blob.name
                chunks = loader.load_and_split()
                if 'https://' or 'http://' in blob.name:
                    for ele in chunks:
                        ele.metadata['source'] = url_part
                        ele.metadata['documents'] = url_part
                else:
                    for ele in chunks:
                        ele.metadata['source'] = s_url
                        ele.metadata['documents'] = file_name

                print(f"Number of chunks :: {len(chunks)}")
                if len(chunks) == 0:
                    tot_fail += 1
                    # Failed files list
                    session['failed_files'].append(file_name)
                    jsonify({"message": f"No text available in {file_name} this file."})

                final_chunks.extend(chunks)

                tot_suc += 1
                if file_name not in session['failed_files']:
                    session['progress_files'].append(file_name)
                for file_name in session['embedding_not_created']:
                    session['embedding_not_created'].remove(file_name)

                # Calculate progress_list
                session['successful_list'] = min(tot_suc - tot_fail, session['total_files_list'])
                session['failed_list'] = min(tot_fail, session['total_files_list'] - session['successful_list'])
                session['progress_list'] = session['total_files_list'] - session['successful_list'] - session[
                    'failed_list']

                if len(chunks) != 0:
                    get_vectorstore(final_chunks, file_name)

                pie_chart_data = create_pie_chart()
                socketio.emit('updatePieChart', pie_chart_data)
                socketio.emit('pending', session['embedding_not_created'])
                socketio.emit('failed', session['failed_files'])
                socketio.emit('success', session['progress_files'])

                if check_stop_flag():
                    write_stop_flag_to_csv(session['login_pin'], 'False')
                    print("Data Load Cancelled")
                    break

                update_bar_chart_from_blob(session, blob_service_client, container_name)
                socketio.emit('progress', {'percentage': 50, 'pin': session['login_pin']})
                # gauge_source_chart_data = gauge_chart_auth()
                # socketio.emit('update_gauge_chart', gauge_source_chart_data)
            socketio.emit('progress', {'percentage': 75, 'pin': session['login_pin']})

        if not check_stop_flag():
            print("Complete")
            elapsed_time = time.time() - start_time
            g.flag = 1  # Set flag to 1 on success1
            logger.info(f"Function update_when_file_delete Data Loaded Successfully in {elapsed_time} seconds")
            return jsonify({"message": "Data loaded successfully"})

        else:
            print("Load Process Stopped")
            elapsed_time = time.time() - start_time
            g.flag = 0  # Set flag to 1 on success1
            logger.info(f"Function update_when_file_delete Data Loaded unsuccessfully in {elapsed_time} seconds")
            return jsonify({"message": "Data loading cancelled!"})

    except Exception as e:
        g.flag = 0  # Set flag to 1 on success1
        logger.error(f"Function update_when_file_delete error", exc_info=True)
        print("update_when_file_delete----->", str(e))
        return jsonify({'message': str(e)})


def generate_word_cloud(word_cloud_text):
    """
    Generates a word cloud image from the provided text and saves it to a file.

    Args:
        word_cloud_text (str): The text to be used for generating the word cloud.

    Side Effects:
        Saves the generated word cloud image to a file in the user's folder.

    Returns:
        None
    """
    # Tokenization
    tokens = word_tokenize(word_cloud_text)

    # Generate WordCloud with specified width and height
    wordcloud = WordCloud(width=1200, height=800,
                          background_color='white',
                          min_font_size=10).generate(' '.join(tokens))
    folder_name = str(session['login_pin'])

    # Save the WordCloud image to a file
    wordcloud.to_file(f'static/login/{folder_name}/wordcloud.png')


def create_or_update_index():
    index_name = f"cognilink-{session['env_map']}-{session['login_pin']}"
    # Create the SearchIndexClient
    client = SearchIndexClient(vector_store_address, AzureKeyCredential(vector_store_password))

    # Define the index schema
    fields = [
        SimpleField(name="id", type=SearchFieldDataType.String, key=True),
        SearchableField(name="file_name", type=SearchFieldDataType.String, filterable=True, sortable=True,
                        searchable=True),
        SearchableField(name="content", type=SearchFieldDataType.String, searchable=True),
        SearchField(name="content_vector", type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
                    searchable=True, vector_search_dimensions=1536, vector_search_profile_name="my-vector-config"),
        SearchableField(name="metadata", type=SearchFieldDataType.String, searchable=True)
    ]
    vector_search = VectorSearch(
        profiles=[VectorSearchProfile(name="my-vector-config", algorithm_configuration_name="my-algorithms-config")],
        algorithms=[HnswAlgorithmConfiguration(name="my-algorithms-config")],
    )

    semantic_config = SemanticConfiguration(
        name="my-semantic-config",
        prioritized_fields=SemanticPrioritizedFields(
            content_fields=[SemanticField(field_name='content'), SemanticField(field_name='file_name')],
            keywords_fields=[SemanticField(field_name='file_name')]
        )
    )
    semantic_search = SemanticSearch(configurations=[semantic_config])

    # Define the scoring profile
    scoring_profile = ScoringProfile(
        name="custom_scoring_profile",
        text_weights=TextWeights(weights={"content": 50, "content_vector": 50}),
    )

    index = SearchIndex(name=index_name, fields=fields,
                        vector_search=vector_search,
                        semantic_search=semantic_search,
                        scoring_profiles=[scoring_profile],
                        )

    # Create or update the index
    try:
        client.create_or_update_index(index)
        logger.info(f"Index {index_name} created for user {str(session['user_name'])}")
        print(f"Index '{index_name}' created or updated successfully.")
    except Exception as e:
        logger.error(f"Index {index_name} not created for user {str(session['user_name'])}", exc_info=True)
        print(f"Index '{index_name}' creation or update failed. Error: {str(e)}")


# Make sure to define session, vector_store_address, and vector_store_password before calling this function


def get_vectorstore(text_chunks, file_name):
    """
    Performs the specified operation on an AzureSearch vector store with the given text chunks.

    Args:
        text_chunks (list): A list of text chunks to be added to the vector store.
        file_name (str): Filename to add a field in vector db

    Side Effects:
        Adds, updates, deletes, or overrides the provided text chunks in the AzureSearch vector store.

    Returns:
        None
    """
    try:
        index_name = f"cognilink-{session['env_map']}-{session['login_pin']}"
        # Create the SearchClient
        search_client = SearchClient(vector_store_address, index_name, AzureKeyCredential(vector_store_password))
        documents = []

        # Replace dots in file_name for vector_id creation
        # Extract the URL part for comparison
        if 'https://' in file_name or 'http://' in file_name:
            file_name = file_name.split("//")[1]
        else:
            file_name = file_name

        cleaned_filename = re.sub(r'[^\w\s.-]', '', file_name)
        cleaned_filename = re.sub(r'\s+', '-', cleaned_filename)
        sanitized_file_name = re.sub(r'[.:/]', '_', cleaned_filename)

        count = 0
        for i, chunk in enumerate(text_chunks):
            embedding = embeddings.embed_query(
                f'/ {file_name} /' + chunk.page_content + f' / {file_name.split(".")[0]}')
            metadata_str = json.dumps(chunk.metadata)
            document = {
                "id": f"{sanitized_file_name}_{str(count)}",
                "file_name": file_name,
                "content": chunk.page_content + f' / {file_name}',
                "content_vector": embedding,
                "metadata": metadata_str
            }
            count += 1
            documents.append(document)

        result = search_client.upload_documents(documents=documents)
        # print(f"Uploaded documents: {result}")
        logger.info('Documents uploaded to vector database')

    except Exception as e:
        logger.error("Documents not uploaded to vector database", exc_info=True)
        print('ERROR! ', e)


def delete_documents_from_vectordb(documents_to_delete):
    """
    Deletes documents from the Azure vector store based on the 'documents' metadata attribute.

    Args:
        documents_to_delete (list): A list of document names to be deleted from the vector store.

    Side Effects:
        Removes the specified documents from the AzureSearch vector store.

    Returns:
        Response: JSON response indicating the result of the operation.
    """
    try:
        search_client = SearchClient(
            endpoint=vector_store_address,
            index_name=f"cognilink-{session['env_map']}-{session['login_pin']}",
            credential=AzureKeyCredential(vector_store_password)
        )

        delete_document_ids = []

        for doc_name in documents_to_delete:
            search_results = search_client.search(search_text=doc_name, search_mode=SearchMode.All)

            # Getting filename from the vectorDB
            for result in search_results:
                delete_document_ids.append(result['id'])

        # Prepare actions for batch delete
        actions = [{"@search.action": "delete", "id": doc_id} for doc_id in delete_document_ids]

        # Use the SearchClient to delete documents
        batch_response = search_client.upload_documents(actions)

        if batch_response:
            logger.info('Documents deleted from vector database')
            print({"message": "Documents deleted successfully"})
        else:
            logger.error('Failed to delete some documents from vector database')
            raise Exception("Failed to delete some documents")

    except Exception as e:
        logger.error('Failed to delete some documents from vector database', exc_info=True)
        print({'message': str(e)})


def get_conversation_chain(retriever, source):
    """
    Retrieves a conversation chain for conversational retrieval using Azure models.

    Args:
        retriever: The AzureSearch vector store instance.
        source (str): The source of the context ('myFiles' or 'webInternet').

    Returns:
        function: A function to handle question answering with context check.
    """

    global template
    if source == 'myFiles':
        template = """ Use the content from the uploaded files by the user. Don't use any web/Internet.
                        If you don't know the answer,just say that you don't know, don't try to make up an answer.
                        {context}
                        Question: {question}
                        Helpful Answer: 
                        """

    elif source == 'webInternet':
        template = """ Search from the Web for the answer and you can use your knowledge also.
                        If you don't know the answer,just say that you don't know, don't try to make up an answer. 
                        {context}
                        Question: {question}
                        Helpful Answer: 
                        """
    elif source == 'all':
        template = """Use both the content from the uploaded files also as well as search internet also and then combine or return the best possible answer.
                    If you don't know the answer,just say that you don't know, don't try to make up an answer.
                    {context} 
                    Question: {question}
                    Helpful Answer: 
                    """

    deployment_name = set_model()
    llm = AzureChatOpenAI(azure_deployment=deployment_name)

    custom_question_prompt = PromptTemplate(input_variables=["question", "context"], template=template)

    memory = ConversationBufferMemory(memory_key="chat_history", input_key='question', return_messages=True,
                                      output_key="answer")

    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        memory=memory,
        condense_question_prompt=custom_question_prompt,
        # verbose=True,
        return_source_documents=True,
        combine_docs_chain_kwargs={"prompt": custom_question_prompt}
    )
    g.flag = 1  # Set flag to 1 on success
    logger.info(f"Retrieval content from conversational chain")
    return conversation_chain


def custom_summary(docs, custom_prompt_, chain_type_, word_count):
    """
    Generates custom summaries for the provided documents using Azure models.

    Args:
        docs (list): List of documents to be summarized.
        custom_prompt_ (str): Custom prompt template for summarization.
        chain_type_ (str): Type of summarization chain.
        word_count (int): Desired word count for the summary.

    Returns:
        list: List of custom summaries.
    """
    # Use the custom prompt as it is for summarization
    combine_prompt = PromptTemplate(template=f"{custom_prompt_}:\n{{text}}\nPlease summarize in {word_count} words.",
                                    input_variables=["text"])
    map_prompt = PromptTemplate(template=f"{custom_prompt_} in {word_count} words:\n{{text}}", input_variables=["text"])
    model = session['engine']
    deployment_name = set_model()
    llm = AzureChatOpenAI(azure_deployment=deployment_name, model_name=model, temperature=0.50)

    if chain_type == "map_reduce":
        chain = load_summarize_chain(llm, chain_type=chain_type_,
                                     map_prompt=map_prompt,
                                     combine_prompt=combine_prompt)
    else:
        chain = load_summarize_chain(llm, chain_type=chain_type)

    summaries = chain({"input_documents": docs}, return_only_outputs=True)["output_text"]
    emit('progress', {'percentage': 50, 'pin': session['login_pin']})

    return summaries


# ## End Summarization section --------


def analyze_sentiment_summ(senti_text_summ):
    """
     Analyzes the sentiment of the provided text and emits the sentiment data.

     Args:
         senti_text_summ (str): The text for sentiment analysis.

     Side Effects:
         Emits sentiment data via socketio.

     Returns:
         None
     """

    # Initialize the Vader Sentiment Intensity Analyzer
    sid = SentimentIntensityAnalyzer()

    # Analyze sentiment
    sentiment_scores = sid.polarity_scores(senti_text_summ)
    # print("Sentiment scores:", sentiment_scores)

    # Calculate percentages
    total = sentiment_scores['pos'] + sentiment_scores['neg'] + sentiment_scores['neu']
    senti_positive_summ = sentiment_scores['pos'] / total * 100
    senti_negative_summ = sentiment_scores['neg'] / total * 100
    senti_neutral_summ = sentiment_scores['neu'] / total * 100

    x1 = [senti_positive_summ, senti_negative_summ, senti_neutral_summ]
    y1 = ['Positive', 'Negative', 'Neutral']
    pin = session['login_pin']

    senti = {'values': x1, 'labels': y1, 'pin': pin}
    # print("Emitting sentiment data:", senti)
    g.flag = 1  # Set flag to 1 on success
    logger.info(f"Summary sentiments analyzed")
    socketio.emit('analyze_sentiment_summ', senti)


def analyze_sentiment_q_a(senti_text_q_a_):
    """
    Analyzes the sentiment of a given text using Vader Sentiment Intensity Analyzer.

    Parameters:
        senti_text_q_a_ (str): The input text for sentiment analysis.

    Returns:
        dict: A dictionary containing the percentages of positive, negative, and neutral sentiments.
    """
    # Initialize the Vader Sentiment Intensity Analyzer
    sid = SentimentIntensityAnalyzer()

    # Analyze sentiment
    sentiment_scores = sid.polarity_scores(senti_text_q_a_)

    # Calculate percentages
    # total = sum(abs(score) for score in sentiment_scores.values())
    total = sentiment_scores['pos'] + sentiment_scores['neg'] + sentiment_scores['neu']
    senti_positive_q_a = sentiment_scores['pos'] / total * 100
    senti_negative_q_a = sentiment_scores['neg'] / total * 100
    senti_neutral_q_a = sentiment_scores['neu'] / total * 100

    x1 = [senti_positive_q_a, senti_negative_q_a, senti_neutral_q_a]
    y1 = ['Positive', 'Negative', 'Neutral']
    pin = session['login_pin']
    senti_q_a = {'values': x1, 'labels': y1, 'pin': pin}
    g.flag = 1  # Set flag to 1 on success
    logger.info(f"Ask QnA sentiments analyzed")
    socketio.emit('analyze_sentiment_q_a', senti_q_a)


def create_bar_chart():
    """
    Creates a bar chart data dictionary based on the session's bar chart data.

    Returns:
        dict: A dictionary containing x and y values for the bar chart.
    """
    # update_bar_chart_from_blob(session, blob_service_client, container_name)
    bar_chart = session['bar_chart_ss']
    # print(bar_chart)
    # Check if 'bar_chart_ss' exists in session
    if bar_chart:
        x1 = list(bar_chart.values())  # Values for the bars
        y1 = list(bar_chart.keys())  # Labels for the bars
        pin = session['login_pin']
    else:
        x1 = [0]  # Values for the bars
        y1 = ["PDF"]  # Labels for the bars
        pin = session['login_pin']

    file_bar = {'values': x1, 'labels': y1, 'pin': pin}

    # bar_json_graph = pio.to_json(file_bar)
    return file_bar


def perform_lda___Q_A(chat_history_list_, num_topics=1, n_top_words=5):
    """
    Performs Latent Dirichlet Allocation (LDA) on the provided chat history to extract topics and their top words.

    Args:
        chat_history_list_ (list): A list of strings containing chat history.
        num_topics (int): Number of topics to extract (default is 1).
        n_top_words (int): Number of top words per topic (default is 5).

    Side Effects:
        Emits LDA topics and their top words via socketio.

    Returns:
        None
    """
    lda_topics_q_a = {'keywords': set()}
    # Declare the global variable

    # Tokenization and stop words removal
    stop_words = set(stopwords.words('english'))

    def preprocess_text(text):
        tokens = word_tokenize(text.lower())
        tokens = [token for token in tokens if token not in stop_words]
        return ' '.join(tokens)

    # Preprocess conversation text
    processed_conversation = preprocess_text(chat_history_list_)

    # Create a CountVectorizer to convert text to a matrix of token counts
    vectorizer = CountVectorizer(stop_words='english')
    X = vectorizer.fit_transform([processed_conversation])

    # Fit the LDA model
    lda = LatentDirichletAllocation(n_components=num_topics, random_state=42)
    lda.fit(X)

    # Aggregate all keywords from each topic
    for topic in lda.components_:
        keywords = [vectorizer.get_feature_names_out()[i] for i in topic.argsort()[:-n_top_words - 1:-1]]
        lda_topics_q_a['keywords'].update(keywords)  # Add keywords to the set

    lda_topics_q_a['keywords'] = list(lda_topics_q_a['keywords'])  # Convert set to list before emitting
    lda_topics_q_a['pin'] = session['login_pin']

    g.flag = 1
    logger.info('QnA topics generated.')
    # Return lda_topics_Q_A
    socketio.emit('lda_topics_QA', lda_topics_q_a)


def perform_lda____summ(senti_text_summ, num_topics=1, n_top_words=3):
    """
    Performs Latent Dirichlet Allocation (LDA) on the provided text for sentiment analysis to extract unique keywords and sends them with the user pin.

    Args:
        senti_text_summ (str): The text for sentiment analysis.
        num_topics (int): Number of topics to extract.
        n_top_words (int): Number of top words per topic.

    Side Effects:
        Emits LDA topics and their top words via socketio, along with the user pin.

    Returns:
        None
    """
    lda_topics_summ = {'keywords': set()}  # Use a set to collect unique keywords
    stop_words = set(stopwords.words('english'))

    def preprocess_text(text):
        tokens = word_tokenize(text.lower())
        return ' '.join(token for token in tokens if token not in stop_words)

    processed_conversation = preprocess_text(senti_text_summ)
    vectorizer = CountVectorizer(stop_words='english')
    X = vectorizer.fit_transform([processed_conversation])

    lda = LatentDirichletAllocation(n_components=num_topics, random_state=42)
    lda.fit(X)

    # Aggregate all keywords from each topic
    for topic in lda.components_:
        keywords = [vectorizer.get_feature_names_out()[i] for i in topic.argsort()[:-n_top_words - 1:-1]]
        lda_topics_summ['keywords'].update(keywords)  # Add keywords to the set

    lda_topics_summ['keywords'] = list(lda_topics_summ['keywords'])  # Convert set to list before emitting
    lda_topics_summ['pin'] = session['login_pin']

    g.flag = 1
    logger.info('Summary topics generated.')
    # Emit the dictionary which now includes the pin and keywords
    socketio.emit('lda_topics_summ', lda_topics_summ)


def create_pie_chart():
    """
     Creates data for a pie chart based on the session's file processing statistics.

     Returns:
         dict: A dictionary containing labels, values, percentages, and text for the pie chart segments.
     """
    if session['total_files_list'] != 0:
        total_files_list = session['total_files_list']
        successful_list = session['successful_list']
        progress_list = session['progress_list']
        failed_list = session['failed_list']
        pin = session['login_pin']

    else:
        total_files_list = 1
        successful_list = 0
        progress_list = 0
        failed_list = 0
        pin = session['login_pin']

    # Calculate percentages
    labels = ['Read', 'In Progress', 'Failed']
    values = [successful_list, progress_list, failed_list]

    # Create text for each segment of the pie chart
    percentages = [(value / total_files_list) * 100 for value in values]
    text = [f"{label}: {value} ({percentage:.2f}%)" for label, value, percentage in zip(labels, values, percentages)]

    pie_chart = {"labels": labels, "values": values, "percentages": percentages, "text": text, "pin": pin}

    # pie_chart_data = json.loads(pie_chart)
    return pie_chart


def gauge_chart_auth():
    """
    Generates data for a gauge chart representing authentication success rate.

    Returns:
        dict: A dictionary containing x and y values for the gauge chart.
    """
    if 'total_success_rate' in session and 'over_all_readiness' in session:
        try:
            over_all_readiness = float(session['over_all_readiness'])
            total_success_rate = float(session['total_success_rate'])
            if over_all_readiness != 0:
                success_rate = round((total_success_rate / over_all_readiness) * 100, 2)
            else:
                success_rate = 0
        except (ValueError, TypeError) as e:
            success_rate = 0
            over_all_readiness = 0
    else:
        success_rate = 0
        over_all_readiness = 0

    pin = session.get('login_pin', None)  # Use .get() to avoid KeyError

    gauge_fig = {'x': [success_rate], 'y': [over_all_readiness], 'pin': pin}
    return gauge_fig


def log_out_forall():
    """
    Logs out the user and clears session data, including chat history, bar chart URLs, and other variables.
    """
    global tot_file
    global chat_history_list

    chat_history_list = []
    tot_file = 0
    # Check if session login pin exists
    if 'login_pin' in session:
        # Define the folder path using session login pin
        folder_name = os.path.join('static', 'login', str(session['login_pin']))

        wordcloud_image = os.path.join(folder_name, 'wordcloud.png')

        if os.path.exists(wordcloud_image):
            os.remove(wordcloud_image)
        g.flag = 1
        logger.info(f"User {session['user_name']} logged out successful")
    session.clear()


def generate_followup_question(answer, chat_history, context):
    """
        Generates a follow-up question based on a given answer, chat history, and context.

        Args:
        answer (str): The answer provided by the user in the conversation.
        chat_history (str): The previous conversation history.
        context (str): Additional context or topics from files to consider for generating the follow-up question.

        Returns:
        str: A follow-up question derived strictly from the content of the files and related to a topic not yet covered in the conversation.
             If no new topics can be found in the files, returns: 'Sorry, I couldn't see anything relevant'.
    """

    deployment_name = set_model()
    llm = AzureChatOpenAI(azure_deployment=deployment_name)

    # Define the follow-up prompt template
    follow_up_prompt_template = PromptTemplate(
        input_variables=['answer', 'chat_history', 'context'],
        template=("Given the answer: {answer} and the previous conversation context: {chat_history}, "
                  "Your primary and only goal is to generate a follow-up question based on the topics from the files that have not been discussed yet using the{context}. "
                  "The follow-up question should be in the form 'Do you also want to know this?'. "
                  "Make sure the follow-up question is derived strictly from the content of the files and is related to a topic not yet covered in the conversation. "
                  "If no new topics can be found in the files, respond with: 'Sorry, I couldn't see anything relevant'. "
                  "Do not create any questions that are not directly based on the content of the files."

                  )
    )
    follow_up_chain = LLMChain(llm=llm, prompt=follow_up_prompt_template)
    follow_up_response = follow_up_chain.run(answer=answer, chat_history=chat_history, context=context)
    return follow_up_response


def extract_text_from_image(file_obj, language):
    try:
        start_time = time.time()
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_path = temp_file.name
            temp_file.write(file_obj.read())

        with open(temp_path, "rb") as image_stream:
            # Initiate the OCR process using the read API
            ocr_result = computervision_client.read_in_stream(image_stream, language=language, raw=True)

            # Extract the operation ID from the response
            operation_location = ocr_result.headers["Operation-Location"]
            operation_id = operation_location.split("/")[-1]

            # Poll for the result
            while True:
                result = computervision_client.get_read_result(operation_id)
                if result.status not in ['notStarted', 'running']:
                    break

            # Extract text from the result
            text = ""
            if result.status == OperationStatusCodes.succeeded:
                for page in result.analyze_result.read_results:
                    for line in page.lines:
                        text += line.text + '\n'

            doc = docx.Document()
            doc_para = doc.add_paragraph(text)

            # Save DOCX to a BytesIO object
            doc_output = io.BytesIO()
            doc.save(doc_output)
            doc_output.seek(0)

            f_name = file_obj.filename
            f_name = f_name.split('.')[0]

            f_name = clean_filename(f_name)

            # Upload the PDF file to Azure Blob Storage
            blob_name = f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}/{f_name}.docx"
            blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
            blob_client.upload_blob(doc_output, blob_type="BlockBlob", overwrite=True)

            elapsed_time = time.time() - start_time
            g.flag = 1
            logger.info(f"Extracted text from image in {elapsed_time} seconds")

    except Exception as e:
        g.flag = 0
        logger.error(f"Text not extracted from image", exc_info=True)
        print(f"Text not extracted from image {e}")


def extract_text_from_pdf(file_obj, language):
    try:
        start_time = time.time()
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_path = temp_file.name
            temp_file.write(file_obj.read())

        # Read the file from the local path
        with open(temp_path, "rb") as pdf_file:
            read_operation = computervision_client.read_in_stream(pdf_file, language=language, raw=True)

        # Check if the operation was successful
        if not read_operation or not read_operation.headers:
            raise Exception("Failed to initiate read operation.")

        # Get the operation location (URL with an ID at the end) from the response
        read_operation_location = read_operation.headers["Operation-Location"]
        if not read_operation_location:
            raise Exception("Failed to get operation location.")

        # Grab the ID from the URL
        operation_id = read_operation_location.split("/")[-1]

        # Wait for the operation to complete
        while True:
            result = computervision_client.get_read_result(operation_id)
            if result.status not in [OperationStatusCodes.not_started, OperationStatusCodes.running]:
                break

        # Print the detected text from each page
        text = ''
        if result.status == OperationStatusCodes.succeeded:
            read_results = result.analyze_result.read_results
            for page in read_results:
                for line in page.lines:
                    # print(line.text)
                    text += line.text + '\n'

            # save the text of scanned pdf in container
            doc = docx.Document()
            doc_para = doc.add_paragraph(text)

            # Save DOCX to a BytesIO object
            doc_output = io.BytesIO()
            doc.save(doc_output)
            doc_output.seek(0)

            f_name = file_obj.filename
            f_name = f_name.split('.')[0]

            f_name = clean_filename(f_name)

            # Upload the PDF file to Azure Blob Storage
            blob_name = f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}/{f_name}.docx"
            blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
            blob_client.upload_blob(doc_output, blob_type="BlockBlob", overwrite=True)

            elapsed_time = time.time() - start_time
            g.flag = 1
            logger.info(f"Extracted text from pdf in {elapsed_time} seconds")

        else:
            g.flag = 0
            logger.error(f"Text not extracted from pdf", exc_info=True)
            print("The operation did not succeed.")

    except Exception as e:
        g.flag = 0
        logger.error(f"Text not extracted from pdf", exc_info=True)
        print(f"An error occurred: {e}")


# Function to check if a URL is valid
def is_url_valid(url):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }

    try:
        response = requests.get(url, headers=headers, timeout=10)
        # print("response_URL--->", response)
        # Check if the response status code is 200 (OK)
        time.sleep(0.5)
        if response.status_code == 200:
            return True
        else:
            return False
    except requests.exceptions.RequestException as e:
        print(f"Request failed: {e}")
        return False


@app.route('/checksession')
def check_session():
    if 'login_pin' in session:
        # Session is Valid
        return jsonify({'sessionValid': True}), 200
    else:
        print("session is expired!!!")
        # Session is expired or invalid
        return jsonify({'sessionValid': False}), 401


@socketio.on('request_chat_history')
def handle_request_chat_history(data):
    date_str = data.get('date')

    if date_str:
        selected_date = datetime.datetime.strptime(date_str, '%Y-%m-%d')
        # Filter chat history by date
        chat_histories = ChatHistory.query.filter(
            ChatHistory.login_pin == session['login_pin'],
            db.func.date(ChatHistory.chat_date) == selected_date.date()
        ).all()

        # Filter summary history by date
        summary_histories = SummaryHistory.query.filter(
            SummaryHistory.login_pin == session['login_pin'],
            db.func.date(SummaryHistory.summary_date) == selected_date.date()
        ).all()
    else:
        # Retrieve all chat history
        chat_histories = ChatHistory.query.filter_by(login_pin=session['login_pin']).all()
        summary_histories = SummaryHistory.query.filter_by(login_pin=session['login_pin']).all()

    # Convert chat history objects to a list of dictionaries
    chat_history_data = [
        {
            'question': chat.question,
            'answer': chat.answer,
            'source': chat.source,
            'page_number': chat.page_number,
        }
        for chat in chat_histories]

    # Convert summary history objects to a list of dictionaries
    summary_history_data = [
        {
            'filename': summary.filename,
            'summary': summary.summary,
        }
        for summary in summary_histories]

    emit('chat_history', {'chat_history': chat_history_data[::-1]})
    emit('summary_history', {'summary_history': summary_history_data[::-1]})


@socketio.on('connect')
def handle_connect():
    bar_chart_data = create_bar_chart()
    gauge_source_chart_data = gauge_chart_auth()

    # socketio.emit('update_pie_chart', pie_chart_data)
    socketio.emit('update_bar_chart', bar_chart_data)
    socketio.emit('update_gauge_chart', gauge_source_chart_data)
    socketio.emit('userName', {'userName': session['user_name'], 'pin': session['login_pin']})

    chat_history_from_db = ChatHistory.query.filter_by(login_pin=session['login_pin']).all()
    chat_history = [{"question": chat.question,
                     "answer": chat.answer,
                     "source": chat.source,
                     "page_number": chat.page_number}
                    for chat in chat_history_from_db]
    emit('chat_history', {'chat_history': chat_history[::-1]})

    database_details = DatabaseDetailsSave.query.filter_by(login_pin=session['login_pin']).all()
    database = [{"database_name": chat.db_name, }
                for chat in database_details]
    db_name = set()
    for db in database:
        db_name.add(db['database_name'])
    # Code to fetch draft data from storage
    container_client_ = blob_service_client.get_container_client(container_name)
    blob_list = container_client_.list_blobs(
        name_starts_with=f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}")

    blobs = []
    for blob in blob_list:
        if (blob.name.split('/')[2].endswith('.csv') or
                blob.name.split('/')[2].endswith('.xlsx') or
                blob.name.split('/')[2].endswith('.xlscd') or
                blob.name.split('/')[2].endswith('.xls')):
            blobs.append({
                'name': blob.name.split('/')[-1],
                'url': f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/{blob.name}",
                # 'draftType': draft_type,
                'status': 'Uploaded'  # or other status based on your logic
            })

    emit('updateAnalystTable', {'blobs': blobs, 'database_name': list(db_name)})


@socketio.on('send_data')
def handle_send_data(data):
    min_date = data.get('minDate')
    max_date = data.get('maxDate')

    # Process the received data as needed
    print(f"Received data: Min Date: {min_date}, Max Date: {max_date}")

    # Emit an event to notify clients about the new data
    emit('data_received', {
        'min_date': min_date,
        'max_date': max_date,
        'message': 'Data received successfully!'
    })


@socketio.on('update_value')
def handle_update_value(data):
    global Limit_By_Size
    Limit_By_Size = data.get('value')
    print('Limit By Size(K/Count)', Limit_By_Size)

    # Emit an event to notify clients about the updated value
    emit('size_value_updated', {'value': Limit_By_Size, 'message': 'Value updated successfully'})


@app.route('/popup_form', methods=['POST'])
def popup_form():
    global mb_pop, file_size_bytes

    start_time = time.time()

    if request.method == 'POST':
        files = request.files.getlist('myFile')
        # source_url = request.form.get('Source_URL', '')
        source_url_list = request.form.get('urls', '[]')  # Get URLs from form data
        source_urls = json.loads(source_url_list)  # Convert JSON string to list
        mb_pop = 0  # Initialize mb_pop before the loop

        # Check if neither files nor URLs are provided
        if not files and not source_urls:
            return jsonify({'message': 'No data provided'}), 400
        if files:
            if not len(files):
                return jsonify({'message': 'File not found'}), 400
            print('name of file is', files)

            for file in files:
                file.seek(0, os.SEEK_END)  # Move the cursor to the end of the file
                file_size_bytes = file.tell()  # Get the current cursor position, which is the file size in bytes
                file.seek(0)  # Reset the cursor back to the beginning of the file
                mb_pop += file_size_bytes / (1024 * 1024)

            mb_p = int(mb_pop)  # Move this line here
            session['Limit_By_Size'] = int(request.form.get('sizeValue'))

            if mb_p >= session['Limit_By_Size'] != 0:
                print('Limit By Size(K/Count) file size exceeds')
                return jsonify({'message': 'Limit by size(k/count) file size exceeds'})
                # Convert bytes to megabytes
            session['MB'] += float("{:.2f}".format(mb_pop))

            for file in files:
                # Check if the file is blank
                check = check_file(file)
                error_files = []
                if check:
                    if check_error(str(check)) == 'corruptFile':
                        print(f"The file {file.filename} is corrupt.")
                        error_files.append(f"The file {file.filename} is corrupt.")
                        socketio.emit('uploadError', {'message': error_files})
                        continue
                    else:
                        error_files.append(f"The file {file.filename} is blank.")
                        socketio.emit('uploadError', {'message': error_files})
                        continue

                scan_source = False
                if any(ext in file.filename for ext in ['.png', '.jpg', '.JPG', '.JPEG', '.jpeg', '.pdf']):
                    lang = request.form.get('selected_language', False)
                    print("lang------>", lang)
                    if lang and '.pdf' in file.filename:
                        extract_text_from_pdf(file, lang)
                        scan_source = True
                    elif lang:
                        extract_text_from_image(file, lang)

                if file.filename.endswith('.mp3'):
                    # print("mp3--------->file received", file.filename)
                    folder_name = os.path.join('static', 'login', str(session['login_pin']))
                    # Ensure the folder exists
                    os.makedirs(folder_name, exist_ok=True)

                    # Save the uploaded file to a temporary location
                    filename = secure_filename(clean_filename(file.filename))
                    temp_file_path = os.path.join(folder_name, filename)
                    file.save(temp_file_path)

                    # Upload the original MP3 file to blob storage
                    blob_name_mp3 = f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}/{filename}"
                    blob_client_mp3 = blob_service_client.get_blob_client(container=container_name, blob=blob_name_mp3)
                    with open(temp_file_path, "rb") as mp3_file:
                        blob_client_mp3.upload_blob(mp3_file, blob_type="BlockBlob",
                                                    content_settings=ContentSettings(content_type="audio/mpeg"),
                                                    overwrite=True)

                    loader = AssemblyAIAudioTranscriptLoader(temp_file_path, api_key="5bbe5761b36b4ff885dbd18836c3c723")
                    chunks = loader.load_and_split()

                    file_name_without_extension = filename.rsplit('.', 1)[0]  # Use filename instead of file.file_name
                    temp_pdf_path = os.path.join(folder_name, f"{file_name_without_extension}.pdf")

                    doc = SimpleDocTemplate(temp_pdf_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    flowables = []

                    # Add a single heading for all chunks
                    heading = "<font size=14><b>{}</b></font><br/><br/>".format(file_name_without_extension)
                    flowables.append(Paragraph(heading, style=styles["Normal"]))

                    # Combine all chunks' text into a single string
                    combined_text = ""
                    for i, chunk in enumerate(chunks):
                        combined_text += chunk.page_content + "\n\n"

                    # Add the combined text under the heading
                    ptext = "<font size=12>{}</font>".format(combined_text)
                    paragraph = Paragraph(ptext, style=styles["Normal"])
                    flowables.append(paragraph)

                    doc.build(flowables)
                    file_name = os.path.basename(temp_pdf_path)

                    with open(temp_pdf_path, "rb") as file:
                        content = file.read()
                    blob_name = f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}/{file_name}"
                    blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
                    blob_client.upload_blob(content, blob_type="BlockBlob",
                                            content_settings=ContentSettings(content_type="application/pdf"),
                                            overwrite=True)

                    # Remove the temporary PDF and MP3 files
                    os.remove(temp_pdf_path)
                    os.remove(temp_file_path)  # Remove the MP3 file

                if not scan_source:
                    upload_to_blob(file, session, blob_service_client, container_name)
        else:
            if not source_urls:
                print('No URLs found')
                return jsonify({'message': 'No URLs found'}), 400

            # Validate each URL
            invalid_urls = [url for url in source_urls if not is_url_valid(url)]
            if invalid_urls:
                print("Some URLs are not valid:", invalid_urls)
                return jsonify({'message': f'Invalid URLs found: {", ".join(invalid_urls)}'}), 400

            # Process each valid URL
            for url in source_urls:
                blob_name = f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}/{url}"
                blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
                blob_client.upload_blob(url, blob_type="BlockBlob", overwrite=True)

        update_bar_chart_from_blob(session, blob_service_client, container_name)

        elapsed_time = time.time() - start_time
        g.flag = 1
        logger.info(f'Data Uploaded Successfully in {elapsed_time} seconds')
        return jsonify({'message': 'Data uploaded successfully'}), 200

    else:
        g.flag = 0
        logger.error('Invalid request method')
        return jsonify({'message': 'Invalid request method'}), 405


@socketio.on('run_query')
def run_query(data):
    db_type = data['dbType']
    hostname = data['hostname']
    port = data['port']
    username = data['username']
    password = data['password']
    database_name = data['db_name']
    session['database_name'].append(database_name)
    folder_name_azure = f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}"
    start_time = time.time()

    new_chat = DatabaseDetailsSave(
        login_pin=session['login_pin'],
        database_type=db_type,
        hostname=hostname,
        db_name=database_name,
        port=port,
        username=username,
        password=password,
    )
    db.session.add(new_chat)
    db.session.commit()

    if db_type == 'MySQL':
        try:
            conn_string = f"mysql+pymysql://{username}:{password}@{hostname}:{port}/{database_name}"

            engine = create_engine(conn_string)
            inspector = inspect(engine)

            schema_with_descriptions = []

            for table_name in inspector.get_table_names():
                for column in inspector.get_columns(table_name):
                    column_info = {
                        'Table': table_name,
                        'Column': column['name'],
                        'Type': str(column['type']),
                        'Description': ''  # Add your descriptions here
                    }
                    schema_with_descriptions.append(column_info)

            df = pd.DataFrame(schema_with_descriptions)

            excel_buffer = io.BytesIO()
            df.to_excel(excel_buffer, engine='openpyxl', index=False)
            excel_buffer.seek(0)
            excel_filename = f'{database_name}_schema.xlsx'
            blob_name = f"{folder_name_azure}/{excel_filename}"
            blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
            blob_client.upload_blob(
                excel_buffer,
                blob_type="BlockBlob",
                content_settings=ContentSettings(
                    content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
                overwrite=True
            )
            excel_file_path = f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/cognilink-{session['env_map']}/{session['login_pin']}/{excel_filename}"
            g.flag = 1
            elapsed_time = time.time() - start_time
            logger.info(f'Fetched data from database in {elapsed_time} seconds')
            emit('excel_response',
                 {'message': f'Database {database_name} connected successfully!! ', 'url': excel_file_path})

        except Exception as ex:
            print(ex)
            g.flag = 0
            emit('excel_response', {'message': 'Unable to connect with database'})

    elif db_type == 'MongoDB':
        client = pymongo.MongoClient(f'mongodb://{username}:{password}@{hostname}:{port}/')
        # db = client[database_name]
        # # result = db.command('eval',query)
        #
        # elapsed_time = time.time() - start_time
        # g.flag = 1
        # logger.info(f'Fetched MongoDB Data in {elapsed_time} seconds')
        # emit('query_success', {'message': 'Data fetched and uploaded successfully.', 'result': str(result)})

    elif db_type == 'SQLServer':
        conn_str = f'DRIVER={{ODBC Driver 18 for SQL Server}};SERVER={hostname},{port};DATABASE={database_name};UID={username};PWD={password};TrustServerCertificate=yes;'
        try:
            engine = create_engine(conn_str)
            inspector = inspect(engine)
            print("Information-------------->", inspector)

            schema_with_descriptions = []

            for table_name in inspector.get_table_names():
                for column in inspector.get_columns(table_name):
                    column_info = {
                        'Table': table_name,
                        'Column': column['name'],
                        'Type': str(column['type']),
                        'Description': ''  # Add your descriptions here
                    }
                    schema_with_descriptions.append(column_info)

            df = pd.DataFrame(schema_with_descriptions)
            print("Data Formed--------->", df)

            excel_buffer = io.BytesIO()
            df.to_excel(excel_buffer, engine='openpyxl', index=False)
            excel_buffer.seek(0)
            excel_filename = f'{database_name}_schema.xlsx'
            blob_name = f"{folder_name_azure}/{excel_filename}"
            blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
            blob_client.upload_blob(
                excel_buffer,
                blob_type="BlockBlob",
                content_settings=ContentSettings(
                    content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
                overwrite=True
            )
            upload_time = time.time()
            excel_file_path = f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/cognilink-{session['env_map']}/{session['login_pin']}/{excel_filename}"
            g.flag = 1
            elapsed_time = time.time() - start_time
            logger.info(f'Fetched data from database in {elapsed_time} seconds')
            emit('excel_response',
                 {'message': 'Schema uploaded successfully ', 'url': excel_file_path, 'upload time': upload_time})

        except Exception as ex:
            print(ex)
            g.flag = 0
            emit('excel_response', {'message': 'Unable to connect with database'})

    else:
        g.flag = 0
        logger.error('Unsupported database type or connection error.')
        emit('excel_response', {'error': 'Unsupported database type or connection error'})


@socketio.on('stop_process')
def stop_process(data):
    login_pin = data['login_pin']

    if login_pin == session['login_pin']:
        print('stop', login_pin, session['login_pin'])
        write_stop_flag_to_csv(login_pin, 'True')

        stop_flag = read_stop_flag_from_csv(login_pin)
        print('Stop Flag:', stop_flag)

        socketio.emit('stop_process_flag', {'flag': stop_flag, 'pin': login_pin})
        socketio.emit('button_response', {'message': 'Operation Cancelled', 'pin': session.get('login_pin')})

        g.flag = 1
        logger.info('Process is stopped or interrupted!!')
        return jsonify({'message': 'Process will be stopped'})
    else:
        g.flag = 0
        logger.error('ERROR! Process is not stopped!!')
        return jsonify({'message': 'ERROR! Process not stopped.'})


@app.route('/Cogni_button', methods=['GET'])
def Cogni_button():
    try:
        start_time = time.time()
        g.flag = 1
        logger.info('CogniLink load button pressed')
        write_stop_flag_to_csv(session['login_pin'], 'False')
        socketio.emit('progress', {'percentage': 10, 'pin': session['login_pin']})
        print('Stop Flag Value-------------------------->', check_stop_flag())
        response = update_when_file_delete()
        if check_stop_flag():
            write_stop_flag_to_csv(session['login_pin'], 'False')
            print("Data Load Cancelled")
            socketio.emit('button_response', {'message': 'Data loaded cancelled', 'pin': session['login_pin']})
        else:
            elapsed_time = time.time() - start_time
            g.flag = 1
            logger.info(f"Load Data Function executed in {elapsed_time} seconds")
            socketio.emit('progress', {'percentage': 100, 'pin': session['login_pin']})
            write_stop_flag_to_csv(session['login_pin'], 'False')
            socketio.emit('button_response', {'message': 'Data loaded successfully', 'pin': session['login_pin']})
        return response
    except Exception as e:
        g.flag = 0
        logger.error("Unhandled exception occurred", exc_info=True)
        socketio.emit('button_response', {'message': str(e), 'pin': session['login_pin']})
        return jsonify({'message': str(e)}), 500


@socketio.on('summary_input')
def handle_summary_input(data):
    global text_word_cloud
    session['summary_add'] = []
    start_time = time.time()
    error_messages = []

    try:
        session['summary_word_count'] = data['value']
        custom_p = data.get('summary_que', "")
        session['summary_word_count'] = data['value']

        if int(session['summary_word_count']) == 0:
            emit('summary_response', {'message': 'Summary word count is zero!'})
            return

        else:
            word_count = int(session['summary_word_count'])
            emit('progress', {'percentage': 10, 'pin': session['login_pin']})

        # Initialize the search client
        search_client = SearchClient(
            endpoint=vector_store_address,
            index_name=f"cognilink-{session['env_map']}-{session['login_pin']}",
            credential=AzureKeyCredential(vector_store_password)
        )

        # Function to extract keywords from the prompt
        def extract_keywords(prompt):
            # Using a simple regex to extract words, you can replace this with a more sophisticated method if needed
            keys = re.findall(r'\w+', prompt)
            return [keyword for keyword in keys if keyword.lower() != ['summary', 'on', 'generate']]

        # Perform the search
        if not custom_p or "all" in custom_p.lower():
            results = search_client.search(
                search_text="*",  # Retrieve all files
                select="*",
                include_total_count=True
            )
        else:
            prompt_embeddings = embeddings.embed_query(custom_p)
            # Extract keywords from the prompt
            keywords = extract_keywords(custom_p)
            # Construct a flexible search query using the keywords
            # exact_match_query = " or ".join([f"search.ismatch('{keyword}', 'content')" for keyword in keywords])
            partial_match_query = " or ".join([f"search.ismatch('{keyword}*', 'file_name')" for keyword in keywords])
            # search_query = f"({exact_match_query}) or ({partial_match_query})"
            search_query = f"{partial_match_query}"
            # print(search_query)
            results = search_client.search(
                search_text=prompt_embeddings,
                select="*",
                include_total_count=True,
                filter=search_query
            )

        # Initialize the list for old_structure documents
        filename_to_docs = {}

        # Function to split content into chunks of a specified size
        def split_content(content, size=8000):
            return [content[i:i + size] for i in range(0, len(content), size)]

        # Process each result
        for result in results:
            try:
                # Extract file_name and content
                filename = result['file_name']
                content = result['content']

                # Split the content into chunks of 8000 characters
                content_chunks = split_content(content, size=8000)

                # Create Document objects with metadata for each chunk
                docs = [Document(page_content=chunk) for chunk in content_chunks]

                # Add the document chunks to the corresponding filename key
                if filename in filename_to_docs:
                    filename_to_docs[filename].extend(docs)
                else:
                    filename_to_docs[filename] = docs

            except KeyError as e:
                error_messages.append(f"KeyError: {e}. Result: {result}")
            except Exception as e:
                error_messages.append(f"An unexpected error occurred while processing result: {str(e)}")

        # Convert the dictionary to a list of tuples
        old_structure = [(filename, docs) for filename, docs in filename_to_docs.items()]

        # Generate summaries with the required structure
        summ = []
        counter = 1

        for filename, documents in old_structure:
            try:
                summary_list = custom_summary(documents, custom_p, chain_type, word_count)
                summary_text = '\n'.join(summary_list)

                # Create summary history record
                summary_record = SummaryHistory(
                    login_pin=session['login_pin'],
                    filename=filename,  # Assuming filename represents the question or context
                    summary=summary_text,
                    summary_date=datetime.datetime.now()
                )
                # Add record to session
                db.session.add(summary_record)
                db.session.commit()

                key = f'{filename}--{counter}--'
                summary_dict = {'key': key, 'value': summary_list}
                summ.append(summary_dict)
                counter += 1
                emit('summary_response', summ)
            except Exception as e:
                error_message = f"An unexpected error occurred for file {filename}: {str(e)}"
                file_name = check_error(str(e))
                if file_name == 'policyError':
                    error_messages.append(f"""An unexpected error occurred for file {filename}:
                                            It seems the content of your file violates some policies.
                                            This could be due to certain words or phrases that are not allowed.
                                            For more details, refer to: https://learn.microsoft.com/en-us/azure/ai-services/openai/concepts/content-filter?tabs=warning%2Cuser-prompt%2Cpython-new""")
                else:
                    error_messages.append(error_message)
                logger.error('Summary generation error: ' + error_message, exc_info=True)

        session['summary_add'].extend(summ)

        emit('progress', {'percentage': 75, 'pin': session['login_pin']})

        senti_text_summ = ' '.join(entry['value'] for entry in summ)
        analyze_sentiment_summ(senti_text_summ)

        generate_word_cloud(senti_text_summ)
        perform_lda____summ(senti_text_summ)

        elapsed_time = time.time() - start_time
        g.flag = 1
        logger.info(f'Summary Generated in {elapsed_time} seconds')
        emit('progress', {'percentage': 100, 'pin': session['login_pin']})

    except Exception as e:
        g.flag = 0
        error_message = f"An unexpected error occurred: {str(e)}"
        file_name = check_error(str(e))
        if file_name == 'valueError':
            error_messages.append("""An unexpected error occurred:
                                The data is not loaded. Please load the data to resolve this error!""")
        if file_name == 'policyError':
            error_messages.append(f"""An unexpected error occurred for file {filename}:
                                                    It seems the content of your file violates some policies.
                                                    This could be due to certain words or phrases that are not allowed.
                                                    For more details, refer to: https://learn.microsoft.com/en-us/azure/ai-services/openai/concepts/content-filter?tabs=warning%2Cuser-prompt%2Cpython-new""")
        else:
            error_messages.append(error_message)
        logger.error('Summary generation error: ' + str(error_messages), exc_info=True)

    finally:
        if error_messages:
            emit('summary_response', {'errors': error_messages})
            emit('summary_response', {'message': 'No data load'})
        emit('progress', {'percentage': 100, 'pin': session['login_pin']})


@socketio.on('clear_chat_summ')
def handle_clear_chat_summ(data):
    global text_word_cloud
    try:
        if 'login_pin' in session:
            folder_name = os.path.join('static', 'login', str(session['login_pin']))
            wordcloud_image = os.path.join(folder_name, 'wordcloud.png')
            if os.path.exists(wordcloud_image):
                os.remove(wordcloud_image)
        lda_topics_summ = {'pin': session['login_pin'], 'keywords': []}
        socketio.emit('lda_topics_summ', lda_topics_summ)
        senti = {'pin': session['login_pin']}
        socketio.emit('analyze_sentiment_summ', senti)

        text_word_cloud = ''
        session['summary_word_count'] = 0
        session['lda_topics_summ'] = {}
        session['senti_Positive_summ'] = 0
        session['senti_Negative_summ'] = 0
        session['senti_neutral_summ'] = 0
        session['summary_add'] = []

        g.flag = 1  # Set flag to 1 on success
        logger.info(f"Summary chat cleared")
        emit('clear_chat_response', {'message': 'Summary cleared successfully'})

    except Exception as e:
        g.flag = 0
        logger.info(f"Summary chat not cleared. ERROR!", exc_info=True)
        print('error in Cleared Chat', str(e))
        emit('clear_chat_response', {'message': str(e)})


@socketio.on('ask_question')
def handle_ask_question(data):
    # Update progress to 10%
    emit('progress', {'percentage': 10, 'pin': session['login_pin']})
    global senti_text_q_a
    start_time = time.time()
    try:
        source = data['source']
        question = data['question']

        if source == "webInternet":
            llm = AzureChatOpenAI(azure_deployment="gpt-35-turbo", model_name="gpt-4", temperature=0.50)

            # Web search tool
            wrapper = DuckDuckGoSearchAPIWrapper(max_results=25)
            emit('progress', {'percentage': 25, 'pin': session['login_pin']})
            web_search_tool = DuckDuckGoSearchRun(api_wrapper=wrapper)

            context = web_search_tool.invoke({"query": question})

            # print("context generated---------->",context)
            generate_prompt = PromptTemplate(
                template="""

                <|begin_of_text|>

                <|start_header_id|>system<|end_header_id|> 

                You are an AI assistant for Web surfing, that synthesizes web search results. 
                Strictly use the following pieces of web search context to answer the question. If you don't know the answer, just say that you don't know. 
                keep the answer concise, and in a human friendly way. 
                Only make direct references to material if provided in the context.

                <|eot_id|>

                <|start_header_id|>user<|end_header_id|>

                Question: {question} 
                Web Search Context: {context} 
                Answer: 

                <|eot_id|>

                <|start_header_id|>assistant<|end_header_id|>""",
                input_variables=["question", "context"],
            )

            # Chain
            generate_chain = generate_prompt | llm
            emit('progress', {'percentage': 50, 'pin': session['login_pin']})

            answer = generate_chain.invoke({"question": question, "context": context})
            response_content = answer.content if isinstance(answer, AIMessage) else str(answer)

            chat_history_list = [{"question": question,
                                  "answer": response_content,
                                  "source": "Web/Internet",
                                  "page_number": "N/A",
                                  "date": datetime.datetime.now()}]
            # print("chat_history_list--->", chat_history_list)

            # Save chat history to the database
            for entry in chat_history_list:
                new_chat = ChatHistory(
                    login_pin=session['login_pin'],
                    question=entry['question'],
                    answer=entry['answer'],
                    source=entry['source'],
                    page_number=entry['page_number'],
                    chat_date=entry['date']
                )
                db.session.add(new_chat)
            db.session.commit()
            emit('progress', {'percentage': 75, 'pin': session['login_pin']})

            senti_text_q_a = ' '.join(entry['answer'] for entry in chat_history_list)
            analyze_sentiment_q_a(senti_text_q_a)
            perform_lda___Q_A(senti_text_q_a)

            # Retrieve chat history from the database
            chat_history_from_db = ChatHistory.query.filter_by(login_pin=session['login_pin']).all()
            chat_history = [{"question": chat.question,
                             "answer": chat.answer,
                             "source": chat.source,
                             "page_number": chat.page_number,
                             "index": chat.id}
                            for chat in chat_history_from_db]
            emit('progress', {'percentage': 100, 'pin': session['login_pin']})
            emit('response', {'chat_history': chat_history[::-1], 'follow_up': 'N/A'})


        else:
            question_embeddings = embeddings.embed_query(question)

            search_client = SearchClient(
                endpoint=vector_store_address,
                index_name=f"cognilink-{session['env_map']}-{session['login_pin']}",
                credential=AzureKeyCredential(vector_store_password)
            )

            # Perform the search
            results = search_client.search(search_text=question_embeddings,
                                           select="*",
                                           top=5,
                                           vector_queries=[VectorizedQuery(vector=question_embeddings,
                                                                           k_nearest_neighbors=5,
                                                                           fields="content_vector")]
                                           )
            documents = []
            for result in results:
                doc = {
                    'content': result['content'],
                    'score': result['@search.score']  # Access the relevance score
                }
                documents.append(doc)

            sorted_documents = sorted(documents, key=lambda x: x['score'], reverse=True)
            # print("sorted documents------------>", sorted_documents)

            vector_store: AzureSearch = AzureSearch(
                azure_search_endpoint=vector_store_address,
                azure_search_key=vector_store_password,
                index_name=f"cognilink-{session['env_map']}-{session['login_pin']}",
                embedding_function=embeddings.embed_query,
            )

            # Update progress to 25%
            emit('progress', {'percentage': 25, 'pin': session['login_pin']})

            # Create the conversation chain handler
            retriever = vector_store.as_retriever(sorted_documents=sorted_documents)
            conversation_chain_handler = get_conversation_chain(retriever, source)
            response = conversation_chain_handler(question)
            # print("Conversational result----------->", response)

            # Update progress to 50%
            emit('progress', {'percentage': 50, 'pin': session['login_pin']})
            sorry_phrases = ['Sorry, there is no information available.', 'Sorry', 'I am sorry',
                             "I can't see anyting relevant", "I'm sorry"]

            if (
                    any(response["answer"].startswith(phrase) for phrase in sorry_phrases) or
                    not response.get("source_documents") or
                    (response.get("source_documents") and not response["source_documents"][0])
            ):
                doc_source = ["N|A"]
                doc_page_num = ["N|A"]

            else:
                # Initialize a set to track seen pages and lists for sources and page numbers
                seen_pages = set()
                doc_source = []
                doc_page_num = []
                docx_sources = []
                docx_page = []

                if source == "all":
                    doc_source.append("Web|Internet")
                    doc_page_num.append("N|A")

                # Iterate over source documents in the response
                for doc in response["source_documents"]:
                    source = f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/cognilink-{str(session['env_map'])}/{str(session['login_pin'])}/" + doc.metadata.get(
                        "source", "N/A")
                    page = doc.metadata.get("page", "N|A")

                    # Add the source to the set of sources
                    if source.endswith('.docx') or page == "N|A":
                        docx_sources.append(source)
                        docx_page.append(page)

                    else:
                        # Adjust the page number to start from 1 if it starts from 0
                        if isinstance(page, int) and page == 0:
                            page = 1
                        elif isinstance(page, int):
                            page += 1

                        page_str = str(page)

                        # Add source and page to the lists if the page has not been seen before
                        if page_str != 'N|A' and page_str not in seen_pages:
                            seen_pages.add(page_str)
                            doc_source.append(source)
                            doc_page_num.append(page_str)
                doc_source.extend(list(set(docx_sources)))
                doc_page_num.extend(docx_page)

            # Flatten the lists to ensure each Q&A pair is aligned with the corresponding sources
            final_chat_hist = [(response['chat_history'][i].content if response['chat_history'][i] else "",
                                response['chat_history'][i + 1].content if response['chat_history'][i + 1] else "",
                                ", ".join(doc_source), ", ".join(doc_page_num))
                               for i in range(0, len(response['chat_history']), 2)]

            # Create a list of dictionaries for the chat history
            chat_history_list = [{"question": chat_pair[0],
                                  "answer": chat_pair[1],
                                  "source": chat_pair[2],
                                  "page_number": chat_pair[3],
                                  "date": datetime.datetime.now()}
                                 for chat_pair in final_chat_hist]

            # Generate follow-up question
            chat_history = conversation_chain_handler.memory.chat_memory.messages
            context = "Use the following pieces of context from the provided files only. Do not use any information from the internet to answer the question at the end."
            follow_up_question = generate_followup_question(response['answer'], chat_history, context)
            # print("Followup_question------------->", follow_up_question)

            senti_text_q_a = ' '.join(entry['answer'] for entry in chat_history_list)

            # Update progress to 75%
            emit('progress', {'percentage': 75, 'pin': session['login_pin']})

            analyze_sentiment_q_a(senti_text_q_a)
            perform_lda___Q_A(senti_text_q_a)

            # Save chat history to the database
            for entry in chat_history_list:
                new_chat = ChatHistory(
                    login_pin=session['login_pin'],
                    question=entry['question'],
                    answer=entry['answer'],
                    source=entry['source'],
                    page_number=entry['page_number'],
                    chat_date=entry['date']
                )
                db.session.add(new_chat)
            db.session.commit()

            # Update progress to 100%
            emit('progress', {'percentage': 100, 'pin': session['login_pin']})
            elapsed_time = time.time() - start_time
            g.flag = 1
            logger.info(f'Answer Generated in {elapsed_time} seconds')

            # Retrieve chat history from the database
            chat_history_from_db = ChatHistory.query.filter_by(login_pin=session['login_pin']).all()
            chat_history = [{"question": chat.question,
                             "answer": chat.answer,
                             "source": chat.source,
                             "page_number": chat.page_number,
                             "index": chat.id}
                            for chat in chat_history_from_db]
            emit('response', {'chat_history': chat_history[::-1], 'follow_up': follow_up_question})

    except Exception as e:
        g.flag = 0
        logger.error('Ask CogniLink answer generation error', exc_info=True)
        print("Exception of ask_question:", str(e))
        emit('response', {'message': 'No response generated'})


@socketio.on('clear_chat')
def handle_clear_chat():
    try:
        # Clear the ChatHistory table for the specific login_pin
        # db.session.query(ChatHistory).filter_by(login_pin=session['login_pin']).delete()
        # db.session.commit()
        # sentiment variable for Q_A
        session['lda_topics_Q_A'] = {}
        session['senti_Positive_Q_A'] = 0
        session['senti_Negative_Q_A'] = 0
        session['senti_neutral_Q_A'] = 0
        session['chat_history_qa'] = []
        senti_q_a = {'pin': session['login_pin']}
        socketio.emit('analyze_sentiment_q_a', senti_q_a)
        lda_topics_q_a = {'pin': session['login_pin'], 'keywords': []}
        socketio.emit('lda_topics_QA', lda_topics_q_a)
        g.flag = 1  # Set flag to 1 on success
        logger.info(f"clear_chat for ask_question route")
        emit('chat_cleared', {'message': 'Chat history cleared successfully'})

    except Exception as e:
        g.flag = 0  # Set flag to 0 on failure
        logger.error(f"clear_chat for ask_question route error", exc_info=True)
        emit('chat_cleared', {'message': str(e)})


async def delete_blob_async(blob_name, container_client_):
    try:
        blob_client_ = container_client_.get_blob_client(blob_name)
        await blob_client_.delete_blob()
        logger.info(f"Successfully deleted blob: {blob_name}")
    except Exception as e:
        logger.error(f"Error deleting blob: {blob_name}, {str(e)}")


@app.route("/delete", methods=["DELETE"])
async def delete_files():
    try:
        start_time = time.time()
        data = request.get_json()
        file_names = data.get('file_names', [])
        if not file_names:
            return jsonify({'message': 'No files specified for deletion'}), 400

        delete_documents_from_vectordb(file_names)

        container_client_ = blob_service_client.get_container_client(container_name)
        blobs = container_client_.list_blobs(
            name_starts_with=f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}")

        blob_names_to_delete = [blob.name for blob in blobs if any(file_name in blob.name for file_name in file_names)]

        if not blob_names_to_delete:
            return jsonify({'error': 'None of the specified files were found in blob storage'}), 404

        await asyncio.gather(*(delete_blob_async(blob_name, container_client) for blob_name in blob_names_to_delete))

        update_bar_chart_from_blob(session, blob_service_client, container_name)

        elapsed_time = time.time() - start_time
        g.flag = 1
        logger.info(f"Selected vault files deleted successfully in {elapsed_time} seconds")
        return jsonify({'message': f'Files {file_names} deleted successfully'})

    except Exception as e:
        g.flag = 0
        logger.error("Error in delete route", exc_info=True)
        return jsonify({'error': str(e)}), 500


@socketio.on('table_update')
def table_update(search_term=None):
    try:
        start_time = time.time()
        # Initialize SearchClient
        search_client = SearchClient(
            endpoint=vector_store_address,
            index_name=f"cognilink-{session['env_map']}-{session['login_pin']}",
            credential=AzureKeyCredential(vector_store_password)
        )
        results = search_client.search(search_text="*", select="*", include_total_count=True)
        vector_list = []

        unique_documents = set()

        for result in results:
            embeddings_dict = json.loads(result['metadata'])
            document = embeddings_dict.get('documents')
            if document and document not in unique_documents:
                vector_list.append(document)
                unique_documents.add(document)

        blobs = container_client.list_blobs(
            name_starts_with=f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}")

        # Exclude files from blobs_chart based on criteria
        blobs_chart = container_client.list_blobs(
            name_starts_with=f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}")
        blob_list = [blob for blob in blobs_chart if
                     not (blob.name.lower().endswith('.csv') or blob.name.lower().endswith('.mp3'))]
        new_blob_list_jpg = [blob for blob in blob_list if
                             not (blob.name.lower().endswith('.jpg') or blob.name.lower().endswith(
                                 '.png') or blob.name.lower().endswith('.jpeg'))]

        # Initialize the deleted files list
        deleted_files_list = []
        delete_file = container_client.list_blobs(
            name_starts_with=f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}")

        # Extract names from delete_file
        delete_file_names = {blob.name.split('/')[-1] for blob in delete_file}
        # Compare delete_file_names with new_blob_list and add items that are not in new_blob_names
        new_blob_names = {blob.name.split('/')[-1] for blob in new_blob_list_jpg}
        for file_name in delete_file_names:
            if file_name not in new_blob_names:
                deleted_files_list.append(file_name)
        failed_files = session.get('failed_files', [])
        embedding_not_created = session.get('embedding_not_created', [])
        # Prepare data with updated statuses
        data = []
        for blob in blobs:
            if 'https://' or 'http://' in blob.name:
                name_source = blob.name.split(str(session['login_pin']) + '/')[1]

                # Convert the date to IST
                date = blob['last_modified']
                date_str = convert_to_ist(date)
                # print("date--------->", date_str)

            if blob.name.split('/')[2] != 'draft':
                file_name = blob.name.split('/')[2]
                if file_name and name_source in vector_list:
                    status = 'U | EC'
                elif file_name in deleted_files_list:
                    status = 'U | N/A'
                elif file_name in failed_files:
                    status = 'U | F'
                elif file_name in embedding_not_created:
                    status = 'U | ENC'
                else:
                    status = 'U | ENC'

                # Filter based on search term
                if search_term and search_term.lower() not in file_name.lower():
                    continue

                data.append({
                    'name': file_name,
                    'date': date_str,
                    'source_url': name_source,
                    'url': f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/{blob.name}",
                    'status': status
                })
                socketio.emit('update_vault_table', data)

        # Calculate overall readiness count
        blob_lent = len(new_blob_list_jpg)
        session['over_all_readiness'] = blob_lent

        tot_succ = len(vector_list)
        if tot_succ > blob_lent:
            tot_succ = blob_lent

        session['total_success_rate'] = tot_succ
        gauge_source_chart_data = gauge_chart_auth()
        socketio.emit('update_gauge_chart', gauge_source_chart_data)

        # Emit the data to the socket channel 'updateTable'
        update_bar_chart_from_blob(session, blob_service_client, container_name)
        socketio.emit('updateTable', data)

        elapsed_time = time.time() - start_time
        g.flag = 1  # Set flag to 1 on success
        logger.info(f"table_update route successfully sent data in {elapsed_time} seconds")
        socketio.emit('update_vault_table', data)
        return jsonify(data)

    except Exception as e:
        g.flag = 0  # Set flag to 0 on error
        logger.error(f"table_update route error", exc_info=True)
        return jsonify({'error': str(e)}), 500


@socketio.on('webcrawler_start')
def webcrawler_start(data):
    start_time = time.time()
    url = data['url']
    login_pin = data['login_pin']
    session['login_pin'] = login_pin  # Store the login_pin in session
    write_stop_flag_to_csv(login_pin, 'False')
    # Create a unique folder for each user's downloaded files
    folder_name = os.path.join('static', 'files', str(login_pin))
    os.makedirs(folder_name, exist_ok=True)

    try:
        # print("Received URL:", url)
        session['current_status'] = "Website URL Received"
        socketio.emit('update_status', {'status': session['current_status'], 'pin': login_pin})
        response = requests.get(url)
        pdf_info_list = extract_pdf_info_from_table(response.content)
        total_files = len(pdf_info_list)
        session['files_downloaded'] = 0
        session['progress_percentage'] = 0
        session['current_status'] = "Crawling in progress..."
        socketio.emit('update_status', {'status': session['current_status'], 'pin': login_pin})
        for pdf_name, pdf_link in pdf_info_list:
            if check_stop_flag():
                write_stop_flag_to_csv(session['login_pin'], 'False')
                session['current_status'] = "Crawling stopped!"
                socketio.emit('update_status', {'status': session['current_status'], 'pin': login_pin})
                socketio.emit('update_progress', {
                    'current_status': session['current_status'],
                    'total_files': total_files,
                    'files_downloaded': session['files_downloaded'],
                    'progress_percentage': session['progress_percentage'],
                    'current_file': current_file,
                    'pin': login_pin
                })
                # print("Crawling Cancelled")
                return jsonify({'message': 'Webcrawler stopped!'})

            try:
                pdf_url = pdf_link

                temp_filename = re.sub(r'[^\w\s.-]', '', pdf_name)
                cleaned_filename = re.sub(r'\s+', '-', temp_filename)

                name = cleaned_filename.replace('/', '-')
                if not name:
                    name = os.path.basename(urlparse(pdf_url).path)
                download_pdf(pdf_url, folder_name, name)
                session['files_downloaded'] += 1
                session['progress_percentage'] = int(session['files_downloaded'] / total_files * 100)
                current_file = name
                socketio.emit('update_progress', {
                    'current_status': session['current_status'],
                    'total_files': total_files,
                    'files_downloaded': session['files_downloaded'],
                    'progress_percentage': session['progress_percentage'],
                    'current_file': current_file,
                    'pin': login_pin
                })
            except Exception as e:
                print(f"Error occurred: {e}")

        if not check_stop_flag():
            session['current_status'] = "File downloaded successfully"
            socketio.emit('update_progress', {
                'current_status': session['current_status'],
                'total_files': total_files,
                'files_downloaded': session['files_downloaded'],
                'progress_percentage': session['progress_percentage'],
                'current_file': current_file,
                'pin': login_pin
            })
            elapsed_time = time.time() - start_time
            g.flag = 1
            write_stop_flag_to_csv(login_pin, 'False')
            logger.info(f"Web Crawling done successfully in {elapsed_time} seconds")
            socketio.emit('update_status', {'status': session['current_status'], 'pin': login_pin})
            return jsonify({'message': 'Files downloaded successfully'})

        else:
            session['current_status'] = "File Downloading Stopped"
            socketio.emit('update_progress', {
                'current_status': session['current_status'],
                'total_files': total_files,
                'files_downloaded': session['files_downloaded'],
                'progress_percentage': session['progress_percentage'],
                'current_file': current_file,
                'pin': login_pin
            })
            write_stop_flag_to_csv(login_pin, 'False')
            g.flag = 0
            logger.error("Web Crawling Stopped!")
            socketio.emit('update_status', {'status': session['current_status'], 'pin': login_pin})
            return jsonify({'message': 'Process stopped!'})

    except Exception as e:
        session['current_status'] = "Error occurred"
        socketio.emit('update_status', {'status': session['current_status'], 'pin': login_pin})
        g.flag = 0
        logger.error("Web crawling error", exc_info=True)
        print("Exception of web crawling:", str(e))
        return jsonify({'message': 'URL not found'})


def extract_pdf_info_from_table(html_content):
    """
    Extracts PDF information from an HTML table.
    Args:
        html_content (str): HTML content containing the table.
    Returns:
        list: A list of tuples containing PDF name and link.
    """
    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')
    # Find all table rows
    rows = soup.find_all('tr')
    pdf_info_list = []
    # Iterate over each row
    for row in rows:
        if check_stop_flag():
            write_stop_flag_to_csv(session['login_pin'], 'False')
            # print("Web Crawling Cancelled")
            break
        # Extract data from each row
        cells = row.find_all('td')
        if len(cells) >= 2:
            pdf_name = cells[0].text.strip()
            pdf_link = cells[1].find('a')['href']
            pdf_info_list.append((pdf_name, pdf_link))
    return pdf_info_list


def download_pdf(url, folder_name, filename):
    """
    Downloads a PDF file from a given URL and saves it to a specified folder with a given filename.

    Args:
        url (str): The URL from which to download the PDF.
        folder_name (str): The name of the folder to save the downloaded PDF.
        filename (str): The name to give the downloaded PDF file (without the .pdf extension).

    Returns:
        None

    Side Effects:
        Updates the global variable `current_file` to the name of the file being downloaded.
        Logs the download progress and success.
    """
    global current_file

    response = requests.get(url)
    file_path = os.path.join(folder_name, filename + '.pdf')

    print(f"Downloading: {filename} from {url} to {file_path}")
    current_file = filename  # Update current file being downloaded

    with open(file_path, 'wb') as f:
        f.write(response.content)
    g.flag = 1
    logger.info(f"Route download_progress success")
    print(f"Downloaded: {filename}")


@socketio.on('fetch_pdf_files')
def handle_fetch_pdf_files():
    directory_path = "static/files/" + session.get('login_pin', '')
    pdf_files = [file for file in os.listdir(directory_path) if file.endswith(".pdf")]
    emit('pdf_files', {"pdf_files": pdf_files})


def delete_file(directory_path, file_name):
    """
    Deletes a specified file from a given directory.

    Args:
        directory_path (str): The path to the directory containing the file.
        file_name (str): The name of the file to be deleted.

    Returns:
        bool: True if the file was successfully deleted, False otherwise.

    Side Effects:
        Sets the global variable `g.flag` to 1 on success and 0 on failure.
        Logs the success or failure of the file deletion process.
    """
    try:
        file_path = os.path.join(directory_path, file_name)
        if os.path.exists(file_path):
            os.remove(file_path)
            g.flag = 1  # Set flag to 1 on success1
            logger.info(f"Function delete_file success")
            return True
        else:
            g.flag = 0
            logger.error(f"Function delete_file error File does not exist", exc_info=True)
            print("File does not exist:", file_path)
            return False

    except Exception as e:
        g.flag = 0  # Set flag to 1 on success1
        logger.error(f"Function delete_file error", exc_info=True)
        print("Error occurred while deleting file:", str(e))
        return False


@socketio.on('delete_pdf_file')
def delete_pdf_file(data):
    try:
        file_name = data['fileName']
        login_pin = data['login_pin']

        if login_pin is None:
            return jsonify({'message': 'User session not found'}), 400

        if data.get('deletePopup') == 'deletepopupn3':
            directory_path = "static/files/" + login_pin
            res = delete_file(directory_path, file_name)
            # Delete the file
            if res is True:
                print("file deleted")
                return socketio.emit('delete_response', {'message': 'Successfully file deleted'})
            else:
                return socketio.emit('delete_response', {'message': 'Failed to delete'})
        else:
            # Assuming you have the folder name for Azure stored in `folder_name_azure`
            blob_name = f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}/{file_name}"
            blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)

            file_path = os.path.join("static/files", login_pin, file_name)
            with open(file_path, "rb") as file:
                file_content = file.read()

            # Upload the content to Azure Blob Storage, overwriting the existing blob if it exists
            blob_client.upload_blob(file_content, blob_type="BlockBlob",
                                    content_settings=ContentSettings(content_type="application/pdf"),
                                    overwrite=True)
            directory_path = "static/files/" + login_pin

            res = delete_file(directory_path, file_name)

        # Delete the file
        if res is True:
            # print("file deleted from temp")
            g.flag = 1  # Set flag to 1 on success1
            logger.info(f"Successfully webcrawler File Loaded In Cognilink Application")
            socketio.emit('delete_response', {'message': 'Successfully file loaded in CogniLink application'})
            socketio.emit('update_table_vault', {'message': 'update_table'})
        else:
            g.flag = 0
            logger.error(f"Failed To Loaded In Cognilink Application", exc_info=True)
            socketio.emit('delete_response', {'message': 'Failed to load file in CogniLink application'})
    except Exception as e:
        g.flag = 0  # Set flag to 1 on success1
        logger.error(f"Error in webcrawler File Loaded In Cognilink Application --select_pdf_file route error is::{e}",
                     exc_info=True)
        socketio.emit('delete_response', {'message': 'Error occurred while deleting file: {}'.format(str(e))}), 500


@socketio.on('eda_process')
def handle_eda_process(data):
    global df, png_file
    img_base64 = None
    start_time = time.time()
    try:
        file_url = data.get('fileUrl')
        if file_url:
            g.flag = 1
            logger.info("SocketIO Eda_Process File name received")
            blob_list_eda = blob_service_client.get_container_client(container_name).list_blobs(
                name_starts_with=f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}")
            for blob in blob_list_eda:
                if blob.name in file_url:
                    blob_client_ = container_client.get_blob_client(blob)
                    blob_data = blob_client_.download_blob().readall()
                    data_stream = BytesIO(blob_data)

                    if file_url.endswith('.xlsx') or file_url.endswith('.xls'):
                        df = pd.read_excel(data_stream)
                    elif file_url.endswith('.csv'):
                        df = pd.read_csv(data_stream)
                    else:
                        emit('eda_response', {'message': 'Unsupported file format', 'success': False})
                        return
                    data_load_time = time.time() - start_time
                    g.flag = 1
                    logger.info(f"SocketIO Eda_Process Data Loaded Successfully in {data_load_time} seconds")
                    emit('eda_response', {'message': 'Data loaded successfully. Ask Virtual Analyst!', 'success': True})
                    return

        question = data.get('question')
        if question:
            g.flag = 1
            logger.info("SocketIO Eda_Process Question received.")
            llm = AzureOpenAI(
                deployment_name="gpt-4-0125-preview",
                api_key=main_key,
                azure_endpoint="https://ea-openai.openai.azure.com/",
                api_version="2023-05-15"
            )

            agent = Agent(df, config={
                "llm": llm,
                "save_charts": True,
                "save_logs": False,
                "enable_cache": False,
                "save_charts_path": 'static/files/image',
                "open_charts": False,
                "max_retries": 1
            })

            output = agent.chat(question)
            output_receiving_time = time.time() - start_time
            g.flag = 1
            logger.info(f"SocketIO Eda_Process output received in {output_receiving_time} seconds")

            if isinstance(output, pd.DataFrame):
                output_json = output.to_json(orient='records')
                output_type = 'table'
            elif isinstance(output, (int, float)):
                output_json = json.dumps(output)
                output_type = 'numeric'
            elif isinstance(output, str):
                output_json = json.dumps(output)
                output_type = 'text'
            else:
                output_json = json.dumps(str(output))
                output_type = 'unknown'

            image_path = 'static/files/image/'
            png_file = None

            if os.path.exists(image_path):
                for file_name in os.listdir(image_path):
                    if file_name.endswith('.png'):
                        png_file = os.path.join(image_path, file_name)
                        break

            if png_file:
                with open(png_file, "rb") as img_file:
                    img_data = img_file.read()
                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                output_json = json.dumps(question)
                output_type = 'text'
                os.remove(png_file)

            response = {
                'success': True,
                'message': 'Question processed successfully!',
                'output_any': output_json,
                'output_type': output_type,
                'image': img_base64
            }
            elapsed_time = time.time() - start_time
            g.flag = 1
            logger.info(f"SocketIO Eda_Process response sent in {elapsed_time} seconds")
            emit('eda_response', response)
        else:
            g.flag = 0
            logger.info('No Question Provided!')
            emit('eda_response', {'success': False, 'message': 'No question provided!'})
    except Exception as e:
        g.flag = 0
        logger.error(f"Error in Eda_Process: {e}", exc_info=True)
        emit('eda_response', {'message': f'Error occurred while EDA process: {str(e)}', 'success': False})


@socketio.on('eda_excel_to_json')
def excel_to_json():
    try:
        start_time = time.time()
        file_list = ['schema.json', 'schema.xlsx', 'schema.xlx']
        container_client_ = blob_service_client.get_container_client(container_name)
        blob_list = container_client_.list_blobs()
        for blob in blob_list:
            if blob.name.split('/')[-1].split('_')[-1] in file_list:
                if blob.name.split('/')[-1].split('_')[-1] == 'schema.json':
                    emit('excel_to_json', {'message': 'Schema received'})
                else:
                    if blob['last_modified'] != blob['creation_time']:
                        excel_file_path = f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/cognilink-{session['env_map']}/{session['login_pin']}/{blob.name.split('/')[-1]}"

                        excel_file = pd.read_excel(excel_file_path)
                        json_data = excel_file.to_json(orient='records', indent=4)

                        folder_name_azure = f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}"
                        blob_name = f"{folder_name_azure}/{blob.name.split('/')[-1].split('_schema')[0]}_{'schema.json'}"
                        blob_client_ = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
                        blob_client_.upload_blob(
                            json_data,
                            blob_type="BlockBlob",
                            content_settings=ContentSettings(
                                content_type="application/json"),
                            overwrite=True
                        )
                    elapsed_time = time.time() - start_time
                    g.flag = 1
                    logger.info(f"Excel to JSON file converted in {elapsed_time} seconds")
                    emit('excel_to_json', {'message': 'Json schema created'})
            else:
                g.flag = 1
                logger.info(f"No file found to convert to JSON")
                pass
    except Exception as e:
        g.flag = 0
        logger.error(f"Excel to JSON file conversion error: {e}", exc_info=True)


@socketio.on('eda_db_process')
def question_answer_on_structure_data(data):
    start_time = time.time()
    database_name = data['database_name'].strip()

    database_details = DatabaseDetailsSave.query.filter_by(login_pin=session['login_pin'],
                                                           db_name=database_name).first()
    db_user = database_details.username
    db_password = database_details.password
    db_host = database_details.hostname
    db_name = database_details.db_name

    timestamp = datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    folder_name_azure = f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}"
    file_name = f"query_results_{timestamp}.csv"

    try:
        query_input = data.get('query_input')
        print("question_passed------>", query_input)

        container_client_ = blob_service_client.get_container_client(container_name)
        blob_list = container_client_.list_blobs()

        expected_filename = f'{database_name}_schema.json'
        file_found = False
        for blob in blob_list:
            if blob.name.split('/')[-1].endswith(expected_filename):
                file_found = True

        if file_found:
            schema_json = (
                f"https://{blob_service_client.account_name}.blob.core.windows.net/"
                f"{container_name}/cognilink-{session['env_map']}/"
                f"{session['login_pin']}/{expected_filename}"
            )

            if schema_json:
                json_link = requests.get(schema_json)
                json_schema = json_link.json()
                schema_str = "\n".join([
                    f"Table: {entry['Table']}, Column: {entry['Column']}, Type: {entry['Type']}, Description: {entry['Description']}"
                    for entry in json_schema])

            template = '''Based on the database schema:
                    {schema}
                    write a SQL query that would answer the user question. Return only the SQL query:

                    Question: "{question}"

                    SQL QUERY:
                    '''
            prompt = PromptTemplate.from_template(template)

            def format_prompt(data):
                return prompt.format(schema=data['schema'], question=data['question'])

            llm = AzureChatOpenAI(azure_deployment="gpt-35-turbo", model_name="gpt-4", temperature=0.50)
            write_query_chain = (RunnablePassthrough()
                                 | format_prompt
                                 | llm
                                 )
            query_generated = write_query_chain.invoke({"schema": schema_str, "question": query_input})
            query_to_run = query_generated.content
            print("Query To Run---------->", query_to_run)

            if query_to_run:
                conn = mysql.connector.connect(
                    host=db_host,
                    user=db_user,
                    password=db_password,
                    database=db_name,
                )

                cursor = conn.cursor()
                cursor.execute(query_to_run)

                columns = [desc[0] for desc in cursor.description]
                results = cursor.fetchall()
                cursor.close()
                conn.close()

                answer_prompt = PromptTemplate.from_template(
                    """Given the following  user question,corresponding sql query,and sql result 
                    you have to formulate the sql result in an human friendly way.
                    Question:{question}
                    SQL Query:{query}
                    SQL Result:{result}
                    Answer:"""
                )
                llm_chain = LLMChain(prompt=answer_prompt, llm=llm)
                answer = llm_chain.invoke({"question": query_input, "query": query_to_run, "result": results})
                print("Answer----------------->", answer['text'])

                df = pd.DataFrame(results, columns=columns)
                df_table = df.head(5).to_html(index=False)
                csv_buffer = io.StringIO()
                df.to_csv(csv_buffer, index=False)
                csv_buffer.seek(0)

                blob_name = f"{folder_name_azure}/{file_name}"
                blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
                blob_client.upload_blob(
                    csv_buffer.getvalue(),
                    blob_type="BlockBlob",
                    content_settings=ContentSettings(content_type="text/csv"),
                    overwrite=True
                )

                container_client_ = blob_service_client.get_container_client(container_name)
                blob_list = container_client_.list_blobs()
                for blob in blob_list:
                    if file_name in blob.name:
                        csv_file_url = f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/{blob.name}"
                        g.flag = 1
                        logger.info(f'CSV File URL: {csv_file_url}')  # Log the draft URL
                        emit('eda_db_response', {'output': answer['text'], 'query': query_to_run, 'url': csv_file_url,
                                                 'df_table': df_table})
                    else:
                        g.flag = 0
                        logger.error('CSV file URL not found!')

                elapsed_time = time.time() - start_time
                g.flag = 1
                logger.info(f'Fetched data from database in {elapsed_time} seconds')
                emit('eda_query_success', {'message': 'Data fetched and uploaded successfully.'})

            else:
                g.flag = 0
                logger.error("Error in db connection no sql query generated!!")
                emit('eda_db_response', {'message': "Error in db connection no sql query generated!!"})

        else:
            emit('eda_db_response', {'message': f"No json file found for {database_name}"})
    except Exception as ex:
        print(ex)
        g.flag = 0
        logger.error('Database connection issue. Error in db connection no sql query generated!!', exc_info=True)
        emit('eda_db_response',
             {'message': "Database connection issue. Error in db connection no sql query generated!!"})


@app.route('/query_table_update', methods=['GET'])
def query_table_update():
    try:
        start_time = time.time()
        # Code to fetch draft data from storage
        container_client_ = blob_service_client.get_container_client(container_name)
        blob_list = container_client_.list_blobs(
            name_starts_with=f"cognilink-{str(session['env_map'])}/{str(session['login_pin'])}")

        blobs = []
        for blob in blob_list:
            if (blob.name.split('/')[2].endswith('.csv') or
                    blob.name.split('/')[2].endswith('.xlsx') or
                    blob.name.split('/')[2].endswith('.xlscd') or
                    blob.name.split('/')[2].endswith('.xls') or
                    blob.name.split('/')[2].endswith('.json')):
                blobs.append({
                    'name': blob.name.split('/')[-1],
                    'url': f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/{blob.name}",
                    # 'draftType': draft_type,
                    'status': 'Uploaded'  # or other status based on your logic
                })
            g.flag = 1  # Set flag to 1 on success
            elapsed_time = time.time() - start_time
            g.flag = 1  # Set flag to 1 on success
            logger.info(f"query_table_update route successfully sent data in {elapsed_time} seconds")
        # socketio.emit('updateAnalystTable', blobs)
        return jsonify(blobs)

    except Exception as e:
        g.flag = 0  # Set flag to 0 on error
        logger.error(f"query_table_update route error", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route("/", methods=["GET", "POST"])
def home():
    # print("request.url------>", request.base_url)
    global logger
    role_names = UserRole.query.with_entities(UserRole.name.distinct()).all()

    if request.method == "POST":
        pin = request.form.get('authpin')
        group_user = request.form.get('Grp_usr')
        engine = request.form.get('engine')
        print(group_user, pin, engine)

        role_id_mapping = {'Admin': 1, 'Guest': 2, 'ML Engine': 3}
        role_id = role_id_mapping.get(group_user)
        user_role = UserRole.query.filter_by(role_id=role_id).first()
        user = UserDetails.query.filter_by(role_id=role_id, login_pin=pin, status='Active').first()

        if role_id is not None and user_role and user:
            log_out_forall()
            session.modified = True
            session['logged_in'] = True
            # taking base url
            session['env_map'] = env_mapping_dict.get(request.base_url)
            session['login_pin'] = user.login_pin
            session['user_name'] = f"{user.first_name.title()} {user.last_name.title()}"
            session['engine'] = engine
            session['bar_chart_ss'] = {}
            session['over_all_readiness'] = 0
            session['total_success_rate'] = 0
            session['MB'] = 0.0
            session['total_files_list'] = 0
            session['successful_list'] = 0
            session['failed_list'] = 0
            session['progress_list'] = 0
            session['lda_topics_summ'] = {}
            session['lda_topics_Q_A'] = {}
            session['senti_Positive_summ'] = 0
            session['senti_Negative_summ'] = 0
            session['senti_neutral_summ'] = 0
            session['senti_Positive_Q_A'] = 0
            session['senti_Negative_Q_A'] = 0
            session['senti_neutral_Q_A'] = 0
            session['chat_history_qa'] = []
            session['summary_add'] = []
            # Lists to store progress of files loading
            session['embedding_not_created'] = []
            session['failed_files'] = []
            session['progress_files'] = []
            session['stop_flag'] = []
            session['database_name'] = []

            folder_name = os.path.join('static', 'login', str(session['login_pin']))
            if not os.path.exists(folder_name):
                os.makedirs(folder_name)

            if 'login_pin' in session:
                logger = setup_database_logger(session['login_pin'])
            else:
                g.flag = 0
                logger.error("User Session Not Found!", exc_info=True)

            logger = CustomLoggerAdapter(logger, {'user_id': session['login_pin']})
            create_or_pass_folder(container_client, session)
            folder_files = os.path.join('static', 'files', str(session['login_pin']))
            if not os.path.exists(folder_files):
                os.makedirs(folder_files)

            g.flag = 1
            logger.info(f"User {str(session['login_pin'])} {session['user_name']} logged in successfully")
            update_bar_chart_from_blob(session, blob_service_client, container_name)
            create_or_update_index()

            return jsonify({'redirect': url_for('data_source')})

        g.flag = 0
        flash('Invalid credentials! Please try again.', 'error')
        return jsonify({'redirect': url_for('home')})

    return render_template('login.html', role_names=role_names)


# Route for logout button
@app.route('/logout')
def logout():
    log_out_forall()
    flash('You have been successfully logged out!', 'success')
    return redirect(url_for('home'))


@app.route('/data_source', methods=['GET', 'POST'])
def data_source():
    update_bar_chart_from_blob(session, blob_service_client, container_name)
    return render_template('datasource.html')


@app.route('/Ask_Question', methods=['GET', 'POST'])
def ask():
    return render_template('ask.html')


@app.route("/Summary", methods=['GET', 'POST'])
def summary():
    session['summary_word_count'] = 0
    return render_template('summary.html')


@app.route('/Virtual_Analyst', methods=['GET', 'POST'])
def virtual_analyst():
    return render_template('virtual_analyst.html')


@app.route('/signup')
def signup():
    return render_template('signup.html')


@app.route('/file_manager')
def file_manager():
    return render_template('webcrawl_file_manager.html')


if __name__ == '__main__':
    socketio.run(app, debug=True)
