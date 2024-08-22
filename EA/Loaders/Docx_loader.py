from typing import Iterator
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument

class DocxDocumentLoader(BaseLoader):
    """A document loader specialized for .docx files that returns page numbers."""

    def __init__(self, file_path: str, chars_per_page: int = 2000) -> None:
        """Initialize the loader with a file path and character limit per page.

        Args:
            file_path: The path to the .docx file to load.
            chars_per_page: Number of characters per page (default is 1000).
        """
        self.file_path = file_path
        self.chars_per_page = chars_per_page

    def lazy_load(self) -> Iterator[Document]:
        """Lazy loader that reads a .docx file and yields Document objects with page numbers."""
        docx = DocxDocument(self.file_path)
        full_text = []

        for para in docx.paragraphs:
            full_text.append(para.text)

        text = "\n".join(full_text)
        total_chars = len(text)
        num_pages = (total_chars // self.chars_per_page) + 1

        for page_num in range(num_pages):
            start_index = page_num * self.chars_per_page
            end_index = min(start_index + self.chars_per_page, total_chars)
            page_content = text[start_index:end_index]
            yield Document(
                page_content=page_content,
                metadata={
                    "source": self.file_path,
                    "page": page_num + 1
                }
            )
